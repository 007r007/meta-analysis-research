"""
Paper-01 Word 文档生成脚本
运行方式（CC 本地 Windows）：
    cd E:\Meta-analysis writing project\projects\paper-01\06-write
    pip install python-docx pillow
    python generate_docx.py

输出：paper01_draft_v1.docx（同目录）
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# ── 路径配置（脚本所在目录）──────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DRAFT_FILE   = os.path.join(BASE_DIR, "paper01_draft_v1.md")
TABLES_FILE  = os.path.join(BASE_DIR, "paper01_tables.md")
REFS_FILE    = os.path.join(BASE_DIR, "paper01_references.md")
FIG_DIR      = os.path.join(BASE_DIR, "paper01_figures")
OUT_FILE     = os.path.join(BASE_DIR, "paper01_draft_v1.docx")

FIG_FILES = {
    1: "figure1_prisma.png",
    2: "figure2_year_dist.png",
    3: "figure3_moderator_summary.png",
}

FIG_CAPTIONS = {
    1: ("Figure 1.", "PRISMA 2020 flow diagram depicting the study selection process. "
        "Records identified from four databases (PubMed, PsycINFO, Web of Science, Scopus) "
        "were deduplicated and screened at title/abstract and full-text stages. "
        "Final included studies: n = 56."),
    2: ("Figure 2.", "Distribution of included studies by publication year (N = 56). "
        "The dashed vertical line indicates the median publication year."),
    3: ("Figure 3.", "Summary of outcome directions across moderating factor subgroups. "
        "Each bar represents the proportion of studies classified as Positive (all outcomes "
        "significant), Mixed (partial), or Null (no significant between-group effects) within "
        "each subgroup. Numbers in parentheses indicate study counts (k)."),
}

# ── 工具函数 ─────────────────────────────────────────

def set_font(run, size_pt=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=0, space_after=0,
                first_line_indent=0, line_spacing=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Inches(first_line_indent)
    if line_spacing:
        from docx.shared import Pt as _Pt
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        pPr = para._p.get_or_add_pPr()
        lSpc = _OE('w:spacing')
        lSpc.set(_qn('w:line'), str(int(line_spacing * 240)))
        lSpc.set(_qn('w:lineRule'), 'auto')
        pPr.append(lSpc)

def add_double_spacing(para):
    """全文双倍行距"""
    pPr = para._p.get_or_add_pPr()
    lSpc = OxmlElement('w:spacing')
    lSpc.set(qn('w:line'), '480')   # 240 = single, 480 = double
    lSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lSpc)

def set_three_line_table(table):
    """三线表样式"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_border(cell, position, size_pt, color='000000'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        border = OxmlElement(f'w:{position}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(int(size_pt * 8)))  # size in 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)

    # Remove all borders first (table level)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    rows = table.rows
    n_rows = len(rows)
    for i, row in enumerate(rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    # Header top: thick (1.5pt → sz=12)
    for cell in rows[0].cells:
        set_border(cell, 'top', 1.5)
    # Header bottom: thin (0.75pt → sz=6)
    for cell in rows[0].cells:
        set_border(cell, 'bottom', 0.75)
    # Table bottom: thick (1.5pt → sz=12)
    for cell in rows[-1].cells:
        set_border(cell, 'bottom', 1.5)

def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK_PAGE if False else None)
    para.runs[0].add_break()
    # simpler:
    doc.add_page_break()

# ── 读取文件 ─────────────────────────────────────────

with open(DRAFT_FILE, encoding='utf-8') as f:
    draft_text = f.read()

with open(TABLES_FILE, encoding='utf-8') as f:
    tables_text = f.read()

with open(REFS_FILE, encoding='utf-8') as f:
    refs_text = f.read()

# ── 创建文档 ─────────────────────────────────────────

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
for margin in ['top_margin','bottom_margin','left_margin','right_margin']:
    setattr(section, margin, Cm(2.54))

# 默认样式
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# 页码（右下）
def add_page_numbers(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, 12)

add_page_numbers(section)

# ── 封面 ─────────────────────────────────────────────

p = doc.add_paragraph()
para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_before=72, space_after=12)
add_double_spacing(p)
run = p.add_run("Working Memory Training in Healthy Older Adults: A Systematic Review of "
                "Transfer Effects, Training Parameters, and Moderating Factors")
set_font(run, 16, bold=True)

for line in ["[Author names to be added]",
             "Target journal: Ageing Research Reviews",
             "Date: 2026-04-06",
             "Registration: Not pre-registered"]:
    p = doc.add_paragraph()
    para_format(p, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_double_spacing(p)
    run = p.add_run(line)
    set_font(run, 12)

doc.add_page_break()

# ── 正文解析函数 ──────────────────────────────────────

def add_heading1(doc, text):
    p = doc.add_paragraph()
    para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
    add_double_spacing(p)
    run = p.add_run(text)
    set_font(run, 12, bold=True)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    add_double_spacing(p)
    run = p.add_run(text)
    set_font(run, 12, bold=True, italic=True)
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    add_double_spacing(p)
    run = p.add_run(text)
    set_font(run, 12, italic=True)
    return p

def add_body_para(doc, text, first_indent=True):
    """正文段落，支持*斜体*标记"""
    p = doc.add_paragraph()
    para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_after=0,
                first_line_indent=0.5 if first_indent else 0)
    add_double_spacing(p)
    # 解析 *italic* 标记
    parts = re.split(r'\*(.*?)\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_font(run, 12, italic=(i % 2 == 1))
    return p

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Inches(0.5)
    pf.right_indent = Inches(0.5)
    pf.space_after  = Pt(6)
    add_double_spacing(p)
    run = p.add_run(text)
    set_font(run, 11, italic=True)
    return p

# ── 解析 Markdown 正文 ────────────────────────────────

def parse_markdown_body(doc, text):
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过元数据和分隔线
        if re.match(r'^(#\s+Working Memory|##\s+Abstract|---|\*\*Authors|Target|Date|Registration)', line):
            i += 1
            continue
        if line.startswith('---'):
            i += 1
            continue

        # 一级标题
        if re.match(r'^## \d+\.', line):
            txt = re.sub(r'^##\s+', '', line)
            add_heading1(doc, txt)
            i += 1
            continue

        # 二级标题
        if re.match(r'^### ', line):
            txt = re.sub(r'^###\s+', '', line)
            add_heading2(doc, txt)
            i += 1
            continue

        # 三级标题
        if re.match(r'^#### ', line):
            txt = re.sub(r'^####\s+', '', line)
            add_heading3(doc, txt)
            i += 1
            continue

        # Abstract 特殊处理
        if line.startswith('## Abstract'):
            add_heading1(doc, 'Abstract')
            i += 1
            continue

        # **Bold:** 开头（abstract 的各段）
        if re.match(r'^\*\*\w+.*?\*\*:', line):
            p = doc.add_paragraph()
            para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_after=0, first_line_indent=0.5)
            add_double_spacing(p)
            m = re.match(r'^\*\*(.*?)\*\*:(.*)', line)
            if m:
                run1 = p.add_run(m.group(1) + ': ')
                set_font(run1, 12, bold=True)
                run2 = p.add_run(m.group(2).strip())
                set_font(run2, 12)
            i += 1
            continue

        # Keywords 行
        if line.startswith('**Keywords:**'):
            p = doc.add_paragraph()
            para_format(p, WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
            add_double_spacing(p)
            run1 = p.add_run('Keywords: ')
            set_font(run1, 12, italic=True)
            run2 = p.add_run(re.sub(r'^\*\*Keywords:\*\*\s*', '', line))
            set_font(run2, 12)
            i += 1
            continue

        # Blockquote (> )
        if line.startswith('> '):
            txt = re.sub(r'^>\s*', '', line)
            add_blockquote(doc, txt)
            i += 1
            continue

        # Note 行（> *Note.*）
        if line.startswith('>'):
            txt = re.sub(r'^>\s*', '', line)
            if txt:
                add_blockquote(doc, txt)
            i += 1
            continue

        # 有序/无序列表
        if re.match(r'^\d+\.\s+\*\*', line) or re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.25)
            add_double_spacing(p)
            # 解析加粗
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                set_font(run, 12, bold=(j % 2 == 1))
            i += 1
            continue

        # References 标题
        if line.strip() == '## References':
            doc.add_page_break()
            add_heading1(doc, 'References')
            i += 1
            continue

        # 普通段落
        if line.strip():
            # 清理 Markdown 粗体（保留斜体）
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            add_body_para(doc, clean)

        i += 1

# ── Abstract 单独处理 ─────────────────────────────────

def add_abstract(doc, text):
    add_heading1(doc, 'Abstract')
    # 提取 Abstract 块
    m = re.search(r'## Abstract\n(.*?)(?=\n---|\n## 1\.)', text, re.DOTALL)
    if not m:
        return
    abstract_block = m.group(1).strip()
    for line in abstract_block.split('\n'):
        line = line.strip()
        if not line:
            continue
        m2 = re.match(r'^\*\*(.*?)\*\*:\s*(.*)', line)
        if m2:
            p = doc.add_paragraph()
            para_format(p, first_line_indent=0.5)
            add_double_spacing(p)
            run1 = p.add_run(m2.group(1) + ': ')
            set_font(run1, 12, bold=True)
            run2 = p.add_run(m2.group(2))
            set_font(run2, 12)
        elif line.startswith('**Keywords:**'):
            p = doc.add_paragraph()
            para_format(p, space_after=12)
            add_double_spacing(p)
            run1 = p.add_run('Keywords: ')
            set_font(run1, 12, italic=True)
            run2 = p.add_run(re.sub(r'\*\*Keywords:\*\*\s*', '', line))
            set_font(run2, 12)
        else:
            add_body_para(doc, line)

# ── 写入正文（不含 Abstract 和 References） ───────────

add_abstract(doc, draft_text)
doc.add_page_break()

# 提取正文 1-5 节
body_match = re.search(r'(## 1\. Introduction.*?)(?=## References)', draft_text, re.DOTALL)
if body_match:
    parse_markdown_body(doc, body_match.group(1))

# ── 写入 References ────────────────────────────────────

doc.add_page_break()
add_heading1(doc, 'References')

ref_lines = refs_text.strip().split('\n')
for line in ref_lines:
    line = line.strip()
    if not line or line.startswith('#') or line == '---':
        continue
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent       = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after       = Pt(6)
    # 单倍行距
    pPr = p._p.get_or_add_pPr()
    lSpc = OxmlElement('w:spacing')
    lSpc.set(qn('w:line'), '240')
    lSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lSpc)
    # 解析斜体（期刊名）
    parts = re.split(r'\*(.*?)\*', line)
    for j, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_font(run, 12, italic=(j % 2 == 1))

# ── 写入 Tables ───────────────────────────────────────

def parse_and_add_table(doc, md_table_block, table_num, table_title, table_note):
    """从 Markdown 表格块创建三线表"""
    doc.add_page_break()

    # 表格标题
    p = doc.add_paragraph()
    para_format(p, space_after=3)
    run1 = p.add_run(f'Table {table_num}. ')
    set_font(run1, 12, bold=True)
    run2 = p.add_run(table_title)
    set_font(run2, 12)

    # 解析 Markdown 表格行
    rows = []
    for line in md_table_block.split('\n'):
        line = line.strip()
        if not line or re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        if line.startswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)

    if not rows:
        doc.add_paragraph('[Table data not found]')
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):
                break
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            is_header = (i == 0)
            set_font(run, 10, bold=is_header)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    set_three_line_table(table)

    # 表注
    if table_note:
        p = doc.add_paragraph()
        para_format(p, space_before=3)
        run1 = p.add_run('Note. ')
        set_font(run1, 10, italic=True)
        run2 = p.add_run(table_note)
        set_font(run2, 10)

# Table 1
t1_match = re.search(r'## Table 1\.(.*?)\n((?:\|.*\n)+)', tables_text, re.DOTALL)
if t1_match:
    parse_and_add_table(
        doc,
        t1_match.group(2),
        1,
        "Characteristics of Included Studies (k = 56)",
        "Design: RCT = randomized controlled trial; Quasi-RCT = quasi-randomized; Single = single-group pre–post; CO = crossover. "
        "Type: NB = n-back; SP = span-based. Adp = adaptive training. "
        "Active ctrl = active control group present. FU = follow-up assessed. "
        "Conclusion = overall between-group transfer outcome. RoB = risk of bias rating. NR = not reported."
    )

# Table 2
t2_match = re.search(r'## Table 2\.(.*?)\n((?:\|.*\n)+)', tables_text, re.DOTALL)
if t2_match:
    parse_and_add_table(
        doc,
        t2_match.group(2),
        2,
        "Risk of Bias Assessment (RoB 2.0 for RCTs; ROBINS-I for Non-Randomized Studies)",
        "D1–D5 = RoB 2.0 domains: D1 randomization, D2 deviations from intended interventions, "
        "D3 missing outcome data, D4 outcome measurement, D5 selection of reported results. "
        "Low = low risk; SC = some concerns; High = high risk. "
        "ROBINS-I ratings: Mod = moderate; Ser = serious; Crit = critical."
    )

# Table 3
t3_match = re.search(r'## Table 3\.(.*?)\n((?:\|.*\n)+)', tables_text, re.DOTALL)
if t3_match:
    parse_and_add_table(
        doc,
        t3_match.group(2),
        3,
        "Summary of Evidence by Moderating Factor Subgroups",
        "k = number of studies in subgroup. Positive = all outcomes showed significant between-group advantage. "
        "Null = no significant between-group effects on any transfer outcome. Mixed = partial significant effects. "
        "Evidence Strength: Consistent = >75% same direction; Moderate = 50–75%; Mixed = <50%; Insufficient = k < 3."
    )

# ── 写入 Figures ──────────────────────────────────────

for fig_num in [1, 2, 3]:
    doc.add_page_break()
    fig_path = os.path.join(FIG_DIR, FIG_FILES[fig_num])
    p = doc.add_paragraph()
    para_format(p, WD_ALIGN_PARAGRAPH.CENTER)
    if os.path.exists(fig_path):
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(6.0))
    else:
        run = p.add_run(f'[Figure {fig_num} — file not found: {FIG_FILES[fig_num]}]')
        set_font(run, 12, italic=True)

    # 图注
    p_caption = doc.add_paragraph()
    para_format(p_caption, WD_ALIGN_PARAGRAPH.LEFT, space_before=6)
    label, caption_text = FIG_CAPTIONS[fig_num]
    run1 = p_caption.add_run(label + ' ')
    set_font(run1, 12, bold=True)
    run2 = p_caption.add_run(caption_text)
    set_font(run2, 12)

# ── 保存 ─────────────────────────────────────────────

doc.save(OUT_FILE)
print(f"\n✅ 文档已生成：{OUT_FILE}")
print(f"   页数估算：正文约 30 页 + 3 张表 + 3 张图")
