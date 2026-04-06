import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('06-write/paper01_draft_v1.docx')
all_text = '\n'.join([p.text for p in doc.paragraphs])
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text += '\n' + cell.text

print("="*65)
print("PAPER-01 PRE-SUBMISSION CHECKLIST")
print("(Based on Paper2 review framework: presubmit_check_part1-6.py)")
print("="*65)

# PART 1: Content Completeness
print("\n[PART 1] Content Completeness")
print("-"*65)

print("\n1a. Placeholder check:")
placeholders = ['XXX', 'TODO', 'TBD', '[INSERT', 'PLACEHOLDER', '???', 'forthcoming']
found_any = False
for ph in placeholders:
    if ph.lower() in all_text.lower():
        count = all_text.lower().count(ph.lower())
        print(f"  [WARN] '{ph}': {count} times")
        found_any = True
# Special: [Author names to be added]
if '[Author names to be added]' in all_text:
    print("  [WARN] '[Author names to be added]' on cover page (intentional?)")
    found_any = True
if not found_any:
    print("  [OK] No placeholders found")

print("\n1b. Section structure:")
sections = ['Abstract', 'Introduction', 'Methods', 'Results', 'Discussion', 'Conclusion', 'References']
for s in sections:
    found = any(s in p.text for p in doc.paragraphs)
    print(f"  {'[OK]' if found else '[MISSING]'} {s}")

print(f"\n1c. Tables: {len(doc.tables)} found (expected 3)")
for i, tbl in enumerate(doc.tables):
    print(f"  Table {i+1}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")

# PART 2: Key Numbers
print("\n[PART 2] Key Numbers Verification")
print("-"*65)

key_nums = {
    'Total records (4168)': r'\b4[,.]?168\b',
    'Full-text assessed (445)': r'\b445\b',
    'Included studies (56)': r'\b56\b',
}
for label, pattern in key_nums.items():
    m = re.findall(pattern, all_text, re.IGNORECASE)
    print(f"  {label}: {len(m)} mentions {'[OK]' if m else '[CHECK]'}")

db_check = {'PubMed': r'PubMed', 'PsycINFO': r'PsycINFO', 'Web of Science': r'Web of Science', 'Scopus': r'Scopus'}
print("\n  Databases mentioned:")
for db, pat in db_check.items():
    n = len(re.findall(pat, all_text))
    print(f"    {db}: {n} times {'[OK]' if n>0 else '[MISSING]'}")

# PART 3: Abstract
print("\n[PART 3] Abstract")
print("-"*65)
abstract_text = ""
in_abstract = False
for para in doc.paragraphs:
    if para.text.strip() == "Abstract":
        in_abstract = True
        continue
    if in_abstract:
        if para.text.strip().startswith("1.") or para.text.strip() == "1. Introduction":
            break
        abstract_text += para.text + " "

wc = len(abstract_text.split())
print(f"  Word count: {wc} {'[OK]' if wc<=300 else '[WARN: may exceed limit]'}")
abstract_elements = {
    'Background': any(w in abstract_text.lower() for w in ['working memory', 'older adult', 'cognitive']),
    'Objective': any(w in abstract_text.lower() for w in ['aim', 'objective', 'to systematically']),
    'Methods': any(w in abstract_text.lower() for w in ['search', 'database', 'review']),
    'Results': any(w in abstract_text.lower() for w in ['found', 'showed', 'across', 'result']),
    'Conclusions': any(w in abstract_text.lower() for w in ['suggest', 'conclude', 'evidence', 'implication']),
    'Keywords': 'Keywords' in abstract_text or 'keyword' in abstract_text.lower(),
}
for elem, present in abstract_elements.items():
    print(f"  {'[OK]' if present else '[MISSING]'} {elem}")

# PART 4: Citations & References
print("\n[PART 4] Citations & References")
print("-"*65)

et_al_bad = re.findall(r'et al[^.\s]', all_text)
print(f"  'et al' without period: {len(et_al_bad)} {'[OK]' if not et_al_bad else '[WARN]'}")
if et_al_bad:
    for item in et_al_bad[:5]:
        print(f"    Sample: ...{item}...")

citation_pattern = r'\(([A-Z][a-zA-Z\-]+(?:\s+(?:et al\.|&|and)\s+[A-Z][a-z]+)?)(?:,)?\s+(\d{4}[a-z]?)\)'
citations = re.findall(citation_pattern, all_text)
print(f"  In-text citations: {len(citations)} total, {len(set(citations))} unique")

ref_start = False
ref_entries = []
for para in doc.paragraphs:
    if para.text.strip() == "References":
        ref_start = True
        continue
    if ref_start and para.text.strip() and para.text[0].isupper():
        ref_entries.append(para.text[:60])
print(f"  Reference entries: {len(ref_entries)}")

doi_count = len(re.findall(r'https://doi\.org|doi\.org', all_text))
print(f"  DOI links in references: {doi_count}")

# PART 5: Statistical Symbols
print("\n[PART 5] Statistical Symbols")
print("-"*65)
stat_symbols = {
    'p values': len(re.findall(r'\bp\s*[=<>]', all_text)),
    'k (studies)': len(re.findall(r'\bk\s*=', all_text)),
    'n (sample)': len(re.findall(r'\bn\s*=', all_text)),
    'N (total)': len(re.findall(r'\bN\s*=', all_text)),
}
for sym, count in stat_symbols.items():
    print(f"  {sym}: {count}")
print("  [NOTE] Verify p, k, n, N, d are italic in Word document")

# PART 6: Critical Content
print("\n[PART 6] Critical Content")
print("-"*65)
checks = {
    'PRISMA 2020': bool(re.search(r'PRISMA', all_text)),
    'PROSPERO / registration': bool(re.search(r'PROSPERO|pre.regist|not pre.regist', all_text, re.IGNORECASE)),
    'SWiM / narrative synthesis': bool(re.search(r'SWiM|narrative synthesis', all_text, re.IGNORECASE)),
    'RoB assessment tool cited': bool(re.search(r'RoB 2|RoB2|ROBINS.I|Sterne', all_text, re.IGNORECASE)),
    'Near/far transfer defined': bool(re.search(r'near transfer|far transfer', all_text, re.IGNORECASE)),
    'Limitations section': bool(re.search(r'[Ll]imitation', all_text)),
    'Publication bias mentioned': bool(re.search(r'publication bias', all_text, re.IGNORECASE)),
    'Conflict of interest statement': bool(re.search(r'conflict of interest|competing interest', all_text, re.IGNORECASE)),
    'Funding statement': bool(re.search(r'funding|grant|no funding', all_text, re.IGNORECASE)),
    'Data availability': bool(re.search(r'data avail|supplementary|supporting information', all_text, re.IGNORECASE)),
    'Ethics/IRB': bool(re.search(r'ethics|IRB|institutional review', all_text, re.IGNORECASE)),
}
for item, present in checks.items():
    print(f"  {'[OK]' if present else '[MISSING]'} {item}")

# PART 7: Internal Consistency
print("\n[PART 7] Internal Consistency")
print("-"*65)

count_56 = len(re.findall(r'\b56\b', all_text))
print(f"  '56 studies': {count_56} mentions [verify consistent]")

cnki = len(re.findall(r'CNKI', all_text))
print(f"  'CNKI' remaining: {cnki} {'[OK]' if cnki==0 else '[WARN: should be Scopus]'}")

first_sr = len(re.findall(r'first systematic review', all_text, re.IGNORECASE))
print(f"  'first systematic review': {first_sr} times {'[CHECK duplicates]' if first_sr>2 else ''}")

overstrong = [
    ('converging evidence', len(re.findall(r'converging evidence', all_text, re.IGNORECASE))),
    ('conclusive', len(re.findall(r'conclusive', all_text, re.IGNORECASE))),
    ('clearly demonstrates', len(re.findall(r'clearly demonstrates', all_text, re.IGNORECASE))),
    ('proves', len(re.findall(r'\bproves?\b', all_text, re.IGNORECASE))),
]
print(f"  Over-strong language check:")
for word, count in overstrong:
    flag = '[CHECK]' if count > 0 else '[OK]'
    print(f"    '{word}': {count} {flag}")

# Exclusion reason sum check
print("\n  Full-text exclusion numbers:")
e1 = len(re.findall(r'E1.*?n\s*=\s*12|12.*?E1', all_text))
e2 = len(re.findall(r'E2.*?n\s*=\s*2|2.*?E2', all_text))
e4 = len(re.findall(r'E4.*?n\s*=\s*1|1.*?E4', all_text))
print(f"    E1=12, E2=2, E4=1 -> total=15 (matches 445-56=389... wait: excluded=15, included=56, fulltext=445 -> 15+56=71 != 445)")
print(f"    [CHECK] 445 full-text - 15 excluded - 56 included = 374 unexplained? Re-verify PRISMA math")

# Actually check: screened 4168, excluded 3723, fulltext 445, ft_excluded 15 -> wait: 445-15=430? No: included=56
# Let's compute
screened = 4168
excl_stage12 = 3723
fulltext = 445
ft_excluded = 15  # from figure1 script
included = 56
screen_math = screened - excl_stage12
ft_math = fulltext - ft_excluded
print(f"\n  PRISMA math verification:")
print(f"    Screened({screened}) - Excluded({excl_stage12}) = {screen_math} (should match fulltext={fulltext}) -> {'OK' if screen_math==fulltext else 'MISMATCH'}")
print(f"    Fulltext({fulltext}) - FT_excluded({ft_excluded}) = {ft_math} (should match included={included}) -> {'CHECK' if ft_math!=included else 'OK'}")
print(f"    NOTE: If ft_math({ft_math}) != included({included}), there may be additional unaccounted records")

print("\n" + "="*65)
print("END OF REPORT")
print("="*65)
