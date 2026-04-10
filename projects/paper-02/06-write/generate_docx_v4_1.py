"""
generate_docx_v4.py
Convert 文档_2_paper02_draft_v2.md → paper02_draft_v4_1.docx

Changes from v3:
  - 导师意见结构性修订：新增§4.3调节因素节（含Table 2）、§4.4机制节重写
  - Introduction §1.3重写（Noble 2015锚点）
  - Abstract修改
  - 源文件改为老师指导后完善/文档_2_paper02_draft_v2.md
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir  = os.path.dirname(os.path.abspath(__file__))
md_path     = os.path.join(script_dir, "老师指导后完善", "文档_2_paper02_draft_v2.md")
docx_path   = os.path.join(script_dir, "老师指导后完善", "paper02_draft_v4_1.docx")
figures_dir = os.path.join(script_dir, "figures")
os.makedirs(os.path.join(script_dir, "老师指导后完善"), exist_ok=True)

# Figure metadata: filename, caption text
FIGURES = [
    (
        "figure1_prisma.png",
        "Figure 1. PRISMA 2020 flow diagram illustrating the study selection process. "
        "Records were identified through searches of PubMed/MEDLINE, PsycINFO, "
        "Web of Science, and Scopus (total = 3,097). After deduplication (n = 1,827), "
        "automated keyword pre-screening (Stage 1) and title/abstract screening (Stage 2) "
        "yielded 133 records for full-text retrieval. Full-text assessment of 109 records "
        "resulted in 16 studies meeting all inclusion criteria."
    ),
    (
        "figure2_modality_age_bubble.png",
        "Figure 2. Distribution of included studies by neural modality and child age at "
        "measurement. Bubble size reflects sample size. Studies are plotted at their "
        "reported mean age (or midpoint of age range). Modalities: rsEEG = resting-state "
        "EEG; ERP = event-related potential; fNIRS = functional near-infrared spectroscopy; "
        "fMRI = functional MRI; DTI = diffusion tensor imaging; sMRI = structural MRI."
    ),
    (
        "figure3_effect_direction.png",
        "Figure 3. Effect direction by neural modality. Each bar represents the proportion "
        "of studies within a modality reporting positive, null, or negative associations "
        "between higher parental education and more mature or stronger neural indices. "
        "Numbers inside bars indicate study counts."
    ),
    (
        "figure4_age_timeline.png",
        "Figure 4. Age at neural measurement across included studies, ordered by modality. "
        "Each point represents one study; horizontal bars indicate the full age range sampled. "
        "Shading denotes the 0–8 year target window. Studies extending beyond 8 years "
        "are included where subgroup analyses or primary effects fell within the target range."
    ),
]

# ── helpers ──────────────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size_pt=12):
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)


def _clear_and_set_spacing(pPr, before_twips=0, after_twips=0,
                            line_twips=480, line_rule="auto"):
    """Remove any existing w:spacing then add a fresh one."""
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"),    str(before_twips))
    spacing.set(qn("w:after"),     str(after_twips))
    spacing.set(qn("w:line"),      str(line_twips))
    spacing.set(qn("w:lineRule"),  line_rule)
    pPr.append(spacing)


def set_para_spacing(para, space_before=0, space_after=6, line_spacing=2.0):
    """Double-spaced by default, 6pt after paragraph."""
    pPr = para._p.get_or_add_pPr()
    _clear_and_set_spacing(
        pPr,
        before_twips=int(space_before * 20),
        after_twips =int(space_after  * 20),
        line_twips  =int(line_spacing * 240),
    )


def add_para(doc, text, style="Normal", bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False):
    para = doc.add_paragraph(style=style)
    para.alignment = align
    set_para_spacing(para)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(1.27)
    add_inline(para, text, bold=bold, italic=italic)
    return para


def add_inline(para, text, bold=False, italic=False):
    """Parse **bold**, *italic*, ***bold-italic*** inline markup and add runs.
    Special case: *Note.* at the very start of text → bold+italic run."""
    # Handle leading *Note.* (APA table note marker)
    note_match = re.match(r'^(\*Note\.\*)(.*)', text, re.DOTALL)
    if note_match:
        run = para.add_run("Note.")
        set_run_font(run, bold=True, italic=True, size_pt=10)
        remainder = note_match.group(2)
        # render the rest normally (may contain more inline markup)
        _add_inline_runs(para, remainder, bold=bold, italic=italic, size_pt=10)
        return
    _add_inline_runs(para, text, bold=bold, italic=italic, size_pt=12)


def _add_inline_runs(para, text, bold=False, italic=False, size_pt=12):
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|[^*]+)', re.DOTALL)
    for m in pattern.finditer(text):
        chunk = m.group(0)
        if chunk.startswith("***") and chunk.endswith("***"):
            run = para.add_run(chunk[3:-3])
            set_run_font(run, bold=True, italic=True, size_pt=size_pt)
        elif chunk.startswith("**") and chunk.endswith("**"):
            run = para.add_run(chunk[2:-2])
            set_run_font(run, bold=True, size_pt=size_pt)
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            run = para.add_run(chunk[1:-1])
            set_run_font(run, italic=True, size_pt=size_pt)
        else:
            run = para.add_run(chunk)
            set_run_font(run, bold=bold, italic=italic, size_pt=size_pt)


def parse_table_md(lines, start):
    """Parse a markdown table starting at lines[start]. Returns (rows, next_idx)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if not all(re.match(r"^[-: ]+$", c) for c in cells):
                rows.append(cells)
        elif rows:
            break
        i += 1
    return rows, i


# ── Three-line table helpers ──────────────────────────────────────────────────

def _border_el(val, sz, color="000000", space="0"):
    """Build a w:XXX border element."""
    el = OxmlElement("w:top")  # tag name set by caller
    el.set(qn("w:val"),   val)
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    return el


def _make_border(tag, val="single", sz="12", color="000000"):
    """Return a border element with the given tag name."""
    el = OxmlElement(tag)
    el.set(qn("w:val"),   val)
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None,
                       insideH=None, insideV=None):
    """Apply explicit borders to a table cell's tcPr."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing tcBorders
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for tag, spec in [("w:top", top), ("w:bottom", bottom),
                       ("w:left", left), ("w:right", right),
                       ("w:insideH", insideH), ("w:insideV", insideV)]:
        if spec is not None:
            el = OxmlElement(tag)
            el.set(qn("w:val"),   spec.get("val",   "single"))
            el.set(qn("w:sz"),    spec.get("sz",    "6"))
            el.set(qn("w:space"), spec.get("space", "0"))
            el.set(qn("w:color"), spec.get("color", "000000"))
            tcBorders.append(el)
        else:
            # explicitly set to none so Word doesn't inherit Table Grid border
            el = OxmlElement(tag)
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            tcBorders.append(el)
    tcPr.append(tcBorders)


THICK = {"val": "single", "sz": "12", "color": "000000"}   # 1.5pt (sz in 1/8pt)
THIN  = {"val": "single", "sz": "6",  "color": "000000"}   # 0.75pt
NONE  = None


def add_three_line_table(doc, rows):
    """Add a three-line (booktabs-style) Word table."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    nrows = len(rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"

    # Remove table-level borders entirely (we control at cell level)
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for tag in ["w:top","w:left","w:bottom","w:right","w:insideH","w:insideV"]:
        el = OxmlElement(tag)
        el.set(qn("w:val"), "none"); el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for ri, row_data in enumerate(rows):
        is_header   = (ri == 0)
        is_last_row = (ri == nrows - 1)

        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            para.clear()
            add_inline(para, cell_text, bold=is_header)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            pPr = para._p.get_or_add_pPr()
            _clear_and_set_spacing(pPr, 0, 0, int(1.15 * 240))

            # Determine borders for this cell
            top_border    = THICK if is_header  else NONE
            bottom_border = THIN  if is_header  else (THICK if is_last_row else NONE)
            _set_cell_borders(
                cell,
                top=top_border,
                bottom=bottom_border,
                left=NONE, right=NONE,
                insideH=NONE, insideV=NONE,
            )

    # Space after table
    doc.add_paragraph()


# ── Figure insertion ──────────────────────────────────────────────────────────

def add_figure_page(doc, fig_path, caption_text):
    """Insert one figure on its own page with caption below."""
    doc.add_page_break()
    # Insert image centred, width = text width (~16cm for 2.54cm margins on A4)
    para_img = doc.add_paragraph()
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para_img, space_before=0, space_after=6, line_spacing=1.0)
    run_img = para_img.add_run()
    run_img.add_picture(fig_path, width=Cm(15.5))

    # Caption paragraph
    para_cap = doc.add_paragraph()
    para_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para_cap, space_before=6, space_after=0, line_spacing=1.5)
    add_inline(para_cap, caption_text)
    for run in para_cap.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


# ── Document setup ─────────────────────────────────────────────────────────

doc = Document()

# Page margins: 2.54 cm all sides
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# Default Normal style — set double spacing at style level
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)
_clear_and_set_spacing(
    style_normal.element.get_or_add_pPr(),
    before_twips=0, after_twips=0, line_twips=480,
)

# Heading styles
for level, (name, size, bold) in enumerate([
    ("Heading 1", 14, True),
    ("Heading 2", 13, True),
    ("Heading 3", 12, True),
    ("Heading 4", 12, True),
], start=1):
    try:
        s = doc.styles[f"Heading {level}"]
    except KeyError:
        s = doc.styles.add_style(f"Heading {level}", WD_STYLE_TYPE.PARAGRAPH)
    s.font.name  = "Times New Roman"
    s.font.size  = Pt(size)
    s.font.bold  = bold
    s.font.color.rgb = RGBColor(0, 0, 0)

# ── Parse markdown ────────────────────────────────────────────────────────────

with open(md_path, encoding="utf-8") as f:
    raw = f.read()

lines = raw.split("\n")

skip_meta_lines = 6
i = 0
in_blockquote   = False
blockquote_lines = []

while i < len(lines):
    line    = lines[i]
    stripped = line.strip()

    # Skip empty
    if not stripped:
        i += 1
        continue

    # Skip horizontal rules
    if stripped in ("---", "***", "___"):
        i += 1
        continue

    # Skip metadata lines (lines 0-5)
    if i < skip_meta_lines and (stripped.startswith("**Draft") or
                                  stripped.startswith("**Status") or
                                  stripped.startswith("**Target") or
                                  stripped == "---"):
        i += 1
        continue

    # Skip placeholder lines
    if stripped.startswith("*(To be compiled") or stripped.startswith("*(Methods are largely"):
        i += 1
        continue

    # H1 title
    if stripped.startswith("# ") and not stripped.startswith("## "):
        title_text = stripped[2:]
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(para, space_before=0, space_after=12, line_spacing=1.5)
        run = para.add_run(title_text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        i += 1
        continue

    # H2 — add page break before major sections
    if stripped.startswith("## ") and not stripped.startswith("### "):
        heading = stripped[3:]
        # Page break before Introduction, Methods, Results, Discussion, Declarations, References
        if any(heading.startswith(k) for k in
               ("1.", "2.", "3.", "4.", "Declarations", "References")):
            doc.add_page_break()
        para = doc.add_heading(heading, level=1)
        set_para_spacing(para, space_before=12, space_after=3, line_spacing=1.5)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.font.bold = True
        i += 1
        continue

    # H3
    if stripped.startswith("### ") and not stripped.startswith("#### "):
        heading = stripped[4:]
        para = doc.add_heading(heading, level=2)
        set_para_spacing(para, space_before=10, space_after=2, line_spacing=1.5)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run.font.bold = True
        i += 1
        continue

    # H4
    if stripped.startswith("#### "):
        heading = stripped[5:]
        para = doc.add_heading(heading, level=3)
        set_para_spacing(para, space_before=8, space_after=2, line_spacing=1.5)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True
        i += 1
        continue

    # Abstract section: detect structured labels **Background:** etc.
    abstract_label_match = re.match(r'^\*\*(.*?:)\*\*\s*(.*)', stripped)
    if abstract_label_match and not stripped.startswith("**Draft") \
            and not stripped.startswith("**Status") \
            and not stripped.startswith("**Target"):
        label = abstract_label_match.group(1)   # e.g. "Background:"
        rest  = abstract_label_match.group(2)
        # Only treat as abstract label for known keywords
        known = ("Background:", "Objective:", "Methods:", "Results:",
                 "Conclusions:", "Keywords:")
        if label in known:
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, space_before=0, space_after=6, line_spacing=2.0)
            run1 = p.add_run(label + " ")
            set_run_font(run1, bold=True, size_pt=12)
            if rest:
                _add_inline_runs(p, rest, size_pt=12)
            # Page break after Keywords line
            if label == "Keywords:":
                doc.add_page_break()
            i += 1
            continue

    # Blockquote (search strategy)
    if stripped.startswith("> "):
        if not in_blockquote:
            in_blockquote    = True
            blockquote_lines = []
        blockquote_lines.append(stripped[2:])
        if i + 1 < len(lines) and lines[i+1].strip().startswith("> "):
            i += 1
            continue
        else:
            bq_text = " ".join(blockquote_lines)
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.left_indent  = Cm(1.27)
            para.paragraph_format.right_indent = Cm(1.27)
            set_para_spacing(para, space_before=3, space_after=3, line_spacing=1.15)
            add_inline(para, bq_text, italic=True)
            in_blockquote    = False
            blockquote_lines = []
            i += 1
            continue

    # Table
    if stripped.startswith("|"):
        rows, next_i = parse_table_md(lines, i)
        if rows:
            add_three_line_table(doc, rows)
            i = next_i
            continue

    # Otherwise normal paragraph
    add_para(doc, stripped, first_line_indent=False)
    i += 1


# ── Figure Captions page ──────────────────────────────────────────────────────

doc.add_page_break()
cap_heading = doc.add_paragraph()
cap_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_para_spacing(cap_heading, space_before=0, space_after=12, line_spacing=1.5)
run_h = cap_heading.add_run("Figure Captions")
set_run_font(run_h, bold=True, size_pt=14)

for fig_file, caption_text in FIGURES:
    fig_path = os.path.join(figures_dir, fig_file)
    if not os.path.exists(fig_path):
        # Add placeholder text if PNG not found
        p = doc.add_paragraph()
        set_para_spacing(p)
        add_inline(p, f"[Figure not found: {fig_file}]", italic=True)
        continue
    # Caption on this page
    p = doc.add_paragraph()
    set_para_spacing(p, space_before=0, space_after=6, line_spacing=1.5)
    add_inline(p, caption_text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

# ── Figures (each on its own page) ───────────────────────────────────────────

for fig_file, caption_text in FIGURES:
    fig_path = os.path.join(figures_dir, fig_file)
    if os.path.exists(fig_path):
        add_figure_page(doc, fig_path, caption_text)

# ── Save ─────────────────────────────────────────────────────────────────────

doc.save(docx_path)
print(f"Saved: {docx_path}")
