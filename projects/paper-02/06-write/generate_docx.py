"""
Paper-02 Word 文档生成脚本  (v2 — 完整投稿格式)
运行方式：
    cd E:\Meta-analysis writing project\projects\paper-02\06-write
    python generate_docx.py
输出：paper02_draft_v1.docx（同目录）

格式标准（Developmental Cognitive Neuroscience / APA）：
- 独立标题页
- 结构式摘要（Background/Objective 等加粗标签）
- 正文 Times New Roman 12pt 双倍行距
- 标题层级：H1 粗体，H2 粗体斜体，H3 斜体
- 三线表（无竖线，顶线1.5pt / 表头底线0.75pt / 底线1.5pt）
- 四张图各独立页，底部图注（Figure N. 加粗 + 正文字体）
- References 悬挂缩进，斜体期刊名
- Supplementary Table S1 单独页
- 正文内 [INSERT TABLE 1 ABOUT HERE] / [INSERT FIGURE N ABOUT HERE] 占位符
- 页码（页脚右对齐，从 1 开始）
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── 路径配置 ──────────────────────────────────────────────────────────────────
BASE     = os.path.dirname(os.path.abspath(__file__))
MD_FILE  = os.path.join(BASE, "文档_2_paper02_draft_v1.md")
NOS_FILE = os.path.join(BASE, "文档_3_Supplementary_TableS1_NOS.md")
FIG_DIR  = os.path.join(BASE, "figures")
OUT_FILE = os.path.join(BASE, "paper02_draft_v1.docx")

FIG_FILES = {
    1: "figure1_prisma.png",
    2: "figure2_modality_age_bubble.png",
    3: "figure3_effect_direction.png",
    4: "figure4_age_timeline.png",
}

FIG_CAPTIONS = {
    1: ("Figure 1.", "PRISMA 2020 flow diagram depicting the study selection process. "
        "Records identified from four databases (PubMed/MEDLINE, PsycINFO, Web of Science, Scopus; "
        "total N = 3,097) were deduplicated (n = 1,827) and screened at automated pre-screening "
        "(Stage 1; n = 513 retained), title/abstract (Stage 2; n = 133 retained), and full-text "
        "stages. Of 109 full texts assessed, 93 were excluded (primary reason: parental education "
        "not independently estimable, E2, k = 87). Final included studies: n = 16."),
    2: ("Figure 2.", "Included studies by neural modality and age group at neural measurement "
        "(N = 16). Each bubble represents one study; bubble area is proportional to sample size (N). "
        "Blue = positive association; red = negative association (k = 1). "
        "† Brito & Noble (2020): p = .025, did not survive FDR correction. "
        "Ramphal et al. (2020): neural measurement at neonatal period; longitudinal follow-up to 2 yr. "
        "Lange et al. (2010): IQ as neural-cognitive proxy; brain volume as covariate."),
    3: ("Figure 3.", "Effect direction by neural modality for all included studies (N = 16). "
        "Bars show the count of studies reporting a positive or negative association between "
        "parental education and neural outcome for each modality. All 15 positive associations "
        "survived covariate adjustment for at least one alternative SES indicator."),
    4: ("Figure 4.", "Age range of neural measurement across included studies. Each horizontal bar "
        "spans the age range of neural measurement for one study. Bar color indicates neural "
        "modality (see legend). Dashed border with hatching = negative association (Wienke et al., 2024). "
        "Numbers at bar right = sample size (N). Background shading indicates developmental age groups. "
        "Ramphal (2020): neonatal fMRI scan; bar extended to 24 months (longitudinal follow-up). "
        "Stiver (2015): preterm sample; bar starts at term-equivalent age."),
}

TABLE1_NOTE = (
    "NOS = Newcastle–Oslo Scale total score (range 0–9; ≥7 = high quality, 5–6 = moderate, "
    "≤4 = low). Effect direction classified as positive when higher parental education was "
    "associated with a more mature, stronger, or larger neural index; negative when higher "
    "parental education was associated with reduced amplitude, lower connectivity, or a less "
    "differentiated neural response. "
    "† Original p = .025; association did not survive false-discovery rate (FDR) correction."
)

# ── 工具函数 ──────────────────────────────────────────────────────────────────

def set_font(run, size_pt=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rPr.insert(0, rFonts)

def double_space(para):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    lSpc = OxmlElement("w:spacing")
    lSpc.set(qn("w:line"), "480")   # 480 twips = double
    lSpc.set(qn("w:lineRule"), "auto")
    pPr.append(lSpc)

def set_spacing(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)

def add_page_numbers(section):
    footer = section.footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    for ftype in ("begin", "end"):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), ftype)
        if ftype == "begin":
            run._r.append(fc)
            instr = OxmlElement("w:instrText")
            instr.text = "PAGE"
            run._r.append(instr)
        else:
            run._r.append(fc)
    set_font(run, 12)

# ── 行内Markdown解析（**bold**, *italic*, ***bold-italic***） ──────────────────

def add_inline(para, text, default_bold=False, default_italic=False,
               size_pt=12, allow_bold=True):
    """Parse **bold** / *italic* / ***bold-italic*** and add runs to para."""
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', re.DOTALL)
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            chunk = text[pos:m.start()]
            run = para.add_run(chunk)
            set_font(run, size_pt, bold=default_bold, italic=default_italic)
        raw = m.group(0)
        if raw.startswith("***"):
            inner = raw[3:-3]
            run = para.add_run(inner)
            set_font(run, size_pt, bold=True if allow_bold else default_bold, italic=True)
        elif raw.startswith("**"):
            inner = raw[2:-2]
            run = para.add_run(inner)
            set_font(run, size_pt, bold=True if allow_bold else default_bold, italic=default_italic)
        else:  # single *
            inner = raw[1:-1]
            run = para.add_run(inner)
            set_font(run, size_pt, bold=default_bold, italic=True)
        pos = m.end()
    # trailing plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        set_font(run, size_pt, bold=default_bold, italic=default_italic)

# ── 段落添加函数 ──────────────────────────────────────────────────────────────

def add_heading1(doc, text):
    """Section heading: 12pt bold (e.g. '1. Introduction')"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=12, after=6)
    double_space(p)
    run = p.add_run(text)
    set_font(run, 12, bold=True)
    return p

def add_heading2(doc, text):
    """Subsection heading: 12pt bold italic"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=6, after=3)
    double_space(p)
    run = p.add_run(text)
    set_font(run, 12, bold=True, italic=True)
    return p

def add_heading3(doc, text):
    """Sub-subsection heading: 12pt italic"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=6, after=3)
    double_space(p)
    run = p.add_run(text)
    set_font(run, 12, italic=True)
    return p

def add_body_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=0, after=0)
    double_space(p)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    add_inline(p, text)
    return p

def add_blockquote(doc, text):
    """Indented italic paragraph for search strings."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    set_spacing(p, before=3, after=3)
    double_space(p)
    run = p.add_run(text)
    set_font(run, 10, italic=True)
    return p

def add_placeholder(doc, text):
    """[INSERT ... ABOUT HERE] centered italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=6, after=6)
    double_space(p)
    run = p.add_run(text)
    set_font(run, 12, italic=True)
    return p

def add_ref_para(doc, text):
    """APA reference: hanging indent 0.5", single-spaced, 6pt after."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after       = Pt(6)
    # single line spacing
    pPr = p._p.get_or_add_pPr()
    lSpc = OxmlElement("w:spacing")
    lSpc.set(qn("w:line"), "240")
    lSpc.set(qn("w:lineRule"), "auto")
    pPr.append(lSpc)
    add_inline(p, text)
    return p

# ── 三线表 ────────────────────────────────────────────────────────────────────

def set_three_line_table(table):
    def set_border(cell, position, size_pt, color="000000"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        border = OxmlElement(f"w:{position}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(int(size_pt * 8)))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)

    # Remove all table-level borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Clear all cell borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                tcB.append(el)
            tcPr.append(tcB)

    rows = table.rows
    # Top border of header row: 1.5pt
    for cell in rows[0].cells:
        set_border(cell, "top", 1.5)
    # Bottom border of header row: 0.75pt
    for cell in rows[0].cells:
        set_border(cell, "bottom", 0.75)
    # Bottom border of last row: 1.5pt
    for cell in rows[-1].cells:
        set_border(cell, "bottom", 1.5)

def add_table_from_md(doc, md_lines, table_num, title, note=None, font_size=9):
    """Parse MD table lines and insert as a three-line table with title and note."""
    doc.add_page_break()
    # Title
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=3)
    run1 = p.add_run(f"Table {table_num}. ")
    set_font(run1, 12, bold=True)
    run2 = p.add_run(title)
    set_font(run2, 12)

    # Parse rows
    rows = []
    for line in md_lines:
        line = line.strip()
        if not line or re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

    if not rows:
        doc.add_paragraph("[Table data not found]")
        return

    n_cols = max(len(r) for r in rows)
    rows   = [r + [""] * (n_cols - len(r)) for r in rows]
    table  = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"   # start point; three-line will override

    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        is_hdr = (ri == 0)
        for ci, cell_text in enumerate(row_data):
            if ci >= len(row.cells):
                break
            cell = row.cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, cell_text, default_bold=is_hdr, size_pt=font_size)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

    set_three_line_table(table)

    # Table note
    if note:
        p = doc.add_paragraph()
        set_spacing(p, before=3)
        run1 = p.add_run("Note. ")
        set_font(run1, 10, bold=True, italic=True)
        run2 = p.add_run(note)
        set_font(run2, 10)

# ── Supplementary Table S1 解析 ───────────────────────────────────────────────

def add_supplementary(doc, nos_file):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=12)
    run = p.add_run("Supplementary Material")
    set_font(run, 14, bold=True)

    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=6)
    run2 = p2.add_run("Supplementary Table S1. Newcastle–Ottawa Scale Domain-Level Scores")
    set_font(run2, 12, bold=True)

    try:
        with open(nos_file, encoding="utf-8") as f:
            nos_text = f.read()
    except FileNotFoundError:
        doc.add_paragraph("[Supplementary Table S1 file not found]")
        return

    # Find main table: between first |---| and summary stats section
    lines = nos_text.split("\n")
    in_table = False
    table_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("| Study") and "Selection" in stripped:
            in_table = True
        if in_table:
            if stripped.startswith("| Study") or (stripped.startswith("|") and
                    not re.match(r"^\|[\s\-:|]+\|$", stripped)):
                table_lines.append(stripped)
            elif re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue  # separator
            elif in_table and stripped == "" and table_lines:
                break

    if table_lines:
        rows = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
        n_cols = max(len(r) for r in rows)
        rows = [r + [""] * (n_cols - len(r)) for r in rows]
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for ri, row_data in enumerate(rows):
            is_hdr = (ri == 0)
            for ci, cell_text in enumerate(row_data):
                if ci >= len(table.rows[ri].cells):
                    break
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                para = cell.paragraphs[0]
                add_inline(para, cell_text, default_bold=is_hdr, size_pt=9)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
        set_three_line_table(table)

    # Note
    p = doc.add_paragraph()
    set_spacing(p, before=3)
    run1 = p.add_run("Note. ")
    set_font(run1, 10, bold=True, italic=True)
    run2 = p.add_run(
        "NOS = Newcastle–Ottawa Scale. Selection (0–4): representativeness, exposure ascertainment. "
        "Comparability (0–2): covariate control. Outcome (0–3): assessment quality and follow-up. "
        "High quality: total ≥7; Moderate: 5–6; Low: ≤4. "
        "Shephard et al. (2019) scores reflect corrected study profile (rsEEG; Brazil; adolescent-mother sample)."
    )
    set_font(run2, 10)

# ── 主 Markdown 解析器 ────────────────────────────────────────────────────────

def parse_main_body(doc, md_text):
    """Parse the full markdown draft and add content to doc in order."""
    lines = md_text.split("\n")
    i = 0
    blockquote_buffer = []

    def flush_blockquote():
        nonlocal blockquote_buffer
        if blockquote_buffer:
            add_blockquote(doc, " ".join(blockquote_buffer))
            blockquote_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank
        if not stripped:
            flush_blockquote()
            i += 1
            continue

        # Skip horizontal rules
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # Skip metadata header lines (first 6 lines)
        if i < 6 and (stripped.startswith("**Draft") or stripped.startswith("**Status") or
                      stripped.startswith("**Target")):
            i += 1
            continue

        # Skip internal notes
        if stripped.startswith("*(Methods are largely") or stripped.startswith("*(To be compiled"):
            i += 1
            continue

        # ── Blockquote ──
        if stripped.startswith("> "):
            flush_blockquote()
            bq_text = stripped[2:]
            # collect multi-line blockquote
            blockquote_buffer.append(bq_text)
            # peek ahead
            while i + 1 < len(lines) and lines[i+1].strip().startswith("> "):
                i += 1
                blockquote_buffer.append(lines[i].strip()[2:])
            flush_blockquote()
            i += 1
            continue

        # ── Document title (H1) ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            flush_blockquote()
            # Title page is added separately; skip
            i += 1
            continue

        # ── H2: major sections ──
        if re.match(r"^## ", stripped) and not re.match(r"^### ", stripped):
            flush_blockquote()
            heading = re.sub(r"^##\s+", "", stripped)

            # Abstract
            if heading == "Abstract":
                add_heading1(doc, "Abstract")
                i += 1
                # Parse abstract structured paragraphs
                while i < len(lines):
                    aline = lines[i].strip()
                    if not aline or aline == "---":
                        if aline == "---":
                            i += 1
                            break
                        i += 1
                        continue
                    if re.match(r"^## ", aline):
                        break
                    # **Label:** text  (colon inside bold markers)
                    m = re.match(r"^\*\*(.*?:)\*\*\s*(.*)", aline)
                    if m:
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Inches(0.5)
                        set_spacing(p)
                        double_space(p)
                        run1 = p.add_run(m.group(1) + " ")
                        set_font(run1, 12, bold=True)
                        add_inline(p, m.group(2))
                    elif aline.startswith("**Keywords:**"):
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Inches(0.5)
                        set_spacing(p, after=12)
                        double_space(p)
                        run1 = p.add_run("Keywords: ")
                        set_font(run1, 12, italic=True)
                        kw = re.sub(r"^\*\*Keywords:\*\*\s*", "", aline)
                        run2 = p.add_run(kw)
                        set_font(run2, 12)
                    i += 1
                doc.add_page_break()
                continue

            # References
            if heading == "References":
                doc.add_page_break()
                add_heading1(doc, "References")
                i += 1
                while i < len(lines):
                    rline = lines[i].strip()
                    if not rline or rline == "---" or rline.startswith("*(To be compiled"):
                        i += 1
                        if not rline:
                            i -= 1
                        if rline == "---":
                            break
                        i += 1
                        continue
                    if re.match(r"^## ", rline):
                        break
                    if rline:
                        add_ref_para(doc, rline)
                    i += 1
                continue

            # Regular section heading with number
            doc.add_page_break()
            add_heading1(doc, heading)
            i += 1
            continue

        # ── H3 ──
        if re.match(r"^### ", stripped):
            flush_blockquote()
            heading = re.sub(r"^###\s+", "", stripped)
            add_heading2(doc, heading)
            i += 1
            continue

        # ── H4 ──
        if re.match(r"^#### ", stripped):
            flush_blockquote()
            heading = re.sub(r"^####\s+", "", stripped)
            add_heading3(doc, heading)
            i += 1
            continue

        # ── Markdown table ──
        if stripped.startswith("|"):
            flush_blockquote()
            # Collect all table lines
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            # Parse rows
            rows = []
            for tline in tbl_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tline):
                    continue
                cells = [c.strip() for c in tline.strip("|").split("|")]
                rows.append(cells)
            if rows:
                n_cols = max(len(r) for r in rows)
                rows = [r + [""] * (n_cols - len(r)) for r in rows]
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for ri, row_data in enumerate(rows):
                    is_hdr = (ri == 0)
                    for ci, cell_text in enumerate(row_data):
                        if ci >= len(table.rows[ri].cells):
                            break
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_inline(para, cell_text, default_bold=is_hdr, size_pt=9)
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after  = Pt(2)
                set_three_line_table(table)
            continue

        # ── Table note (*Note.*) ──
        if stripped.startswith("*Note.*") or stripped.startswith("*Note."):
            flush_blockquote()
            p = doc.add_paragraph()
            set_spacing(p, before=3)
            clean = re.sub(r"^\*Note\.\*\s*", "", stripped)
            run1 = p.add_run("Note. ")
            set_font(run1, 10, bold=True, italic=True)
            add_inline(p, clean, size_pt=10)
            i += 1
            continue

        # ── Normal body paragraph ──
        flush_blockquote()
        add_body_para(doc, stripped)
        i += 1

    flush_blockquote()

# ── 图片页 ────────────────────────────────────────────────────────────────────

def add_figure_page(doc, fig_num):
    doc.add_page_break()
    # Figure image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=6)
    fig_path = os.path.join(FIG_DIR, FIG_FILES[fig_num])
    if os.path.exists(fig_path):
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(6.0))
    else:
        run = p.add_run(f"[Figure {fig_num} — file not found: {FIG_FILES[fig_num]}]")
        set_font(run, 12, italic=True)
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p_cap, before=6, after=0)
    label, caption_body = FIG_CAPTIONS[fig_num]
    run1 = p_cap.add_run(label + " ")
    set_font(run1, 12, bold=True)
    run2 = p_cap.add_run(caption_body)
    set_font(run2, 12)

# ── Figure Captions 列表页 ────────────────────────────────────────────────────

def add_figure_captions_page(doc):
    doc.add_page_break()
    add_heading1(doc, "Figure Captions")
    for fig_num in range(1, 5):
        p = doc.add_paragraph()
        set_spacing(p, before=6, after=0)
        double_space(p)
        label, caption_body = FIG_CAPTIONS[fig_num]
        run1 = p.add_run(label + " ")
        set_font(run1, 12, bold=True)
        run2 = p.add_run(caption_body)
        set_font(run2, 12)

# ── 主程序 ────────────────────────────────────────────────────────────────────

with open(MD_FILE, encoding="utf-8") as f:
    md_text = f.read()

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
    setattr(section, attr, Cm(2.54))

# Default Normal style
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)
# Set double spacing on Normal style so it propagates to all paragraphs
style_pPr = doc.styles["Normal"].element.get_or_add_pPr()
for old in style_pPr.findall(qn("w:spacing")):
    style_pPr.remove(old)
_spc = OxmlElement("w:spacing")
_spc.set(qn("w:line"), "480")
_spc.set(qn("w:lineRule"), "auto")
_spc.set(qn("w:after"), "0")
style_pPr.append(_spc)

# Page numbers
add_page_numbers(section)

# ── 1. 标题页 ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=72, after=12)
double_space(p)
run = p.add_run(
    "Parental Education Level and Early Childhood Neural Development: "
    "A Systematic Review of EEG, ERP, fNIRS, and Neuroimaging Evidence "
    "(Ages 0–8 Years)"
)
set_font(run, 14, bold=True)

for line in [
    "[Author names to be added]",
    "[Institutional affiliations to be added]",
    "Target journal: Developmental Cognitive Neuroscience",
    "Manuscript type: Systematic Review",
    "Date: 2026-04-10",
    "Registration: Not pre-registered (PROSPERO)",
    "Word count: ~9,500 words (excluding abstract, references, tables, figure captions)",
    "Figures: 4  |  Tables: 1  |  Supplementary tables: 1",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=4)
    double_space(p)
    run = p.add_run(line)
    set_font(run, 12)

# Corresponding author placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=24, after=4)
double_space(p)
run = p.add_run("Corresponding author: [Name, email, address to be added]")
set_font(run, 12)

doc.add_page_break()

# ── 2. 正文（Abstract → References） ─────────────────────────────────────────
parse_main_body(doc, md_text)

# ── 3. Figure Captions 列表页 ─────────────────────────────────────────────────
add_figure_captions_page(doc)

# ── 4. 四张图各独立页 ──────────────────────────────────────────────────────────
for n in range(1, 5):
    add_figure_page(doc, n)

# ── 5. Supplementary Table S1 ─────────────────────────────────────────────────
add_supplementary(doc, NOS_FILE)

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
file_kb = os.path.getsize(OUT_FILE) / 1024
print(f"Size: {file_kb:.1f} KB")
