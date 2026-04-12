"""
生成修改说明Word文档：revision_note_to_supervisor.docx
内容：三条核心知识改动 + 四条大论文接口（三线表）+ 大型队列说明 + 一句话摘要
格式：Times New Roman 12pt，双倍行距，三线表，APA风格
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir = os.path.dirname(os.path.abspath(__file__))
out_path   = os.path.join(script_dir, "revision_note_to_supervisor.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def clear_spacing(pPr):
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)

def set_spacing(para, before=0, after=0, line=480, rule="auto"):
    pPr = para._p.get_or_add_pPr()
    clear_spacing(pPr)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),   str(before))
    sp.set(qn("w:after"),    str(after))
    sp.set(qn("w:line"),     str(line))
    sp.set(qn("w:lineRule"), rule)
    pPr.append(sp)

def add_para(doc, text="", bold=False, italic=False, size=12,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0,
             line=480, rule="auto", indent_cm=0):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after, line=line, rule=rule)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_mixed(doc, parts, before=0, after=0, line=480,
              align=WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after, line=line)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

# ── Three-line table helpers ──────────────────────────────────────────────────

THICK = "15840"   # 1.25pt in twentieths-of-a-point... actually use EMUs via w:sz
# word border sz unit = 1/8 pt  → 12 = 1.5pt, 6 = 0.75pt

def _set_table_border(tbl, top_sz="12", bottom_sz="12"):
    """Set table-level top and bottom borders only."""
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    # remove existing tblBorders
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        if side == "top":
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    top_sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
        elif side == "bottom":
            b.set(qn("w:val"),   "single")
            b.set(qn("w:sz"),    bottom_sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
        else:
            b.set(qn("w:val"),   "none")
            b.set(qn("w:sz"),    "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def _set_cell_border(cell, top="none", bottom="none", left="none", right="none",
                     sz_top="6", sz_bottom="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side, val, sz in [
        ("top",    top,    sz_top),
        ("bottom", bottom, sz_bottom),
        ("left",   "none", "0"),
        ("right",  "none", "0"),
        ("insideH","none", "0"),
        ("insideV","none", "0"),
    ]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   val)
        b.set(qn("w:sz"),    sz if val != "none" else "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000" if val != "none" else "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def cell_para(cell, parts, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT,
              before=40, after=40):
    """Replace cell content with a mixed-run paragraph."""
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    tc = cell._tc
    p_elem = OxmlElement("w:p")
    tc.append(p_elem)
    from docx.text.paragraph import Paragraph
    p = Paragraph(p_elem, cell)
    p.alignment = align
    set_spacing(p, before=before, after=after, line=240, rule="auto")
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.18)
    section.right_margin  = Cm(3.18)

# Fix Normal style
normal = doc.styles["Normal"]
nPr    = normal.paragraph_format
nPr.line_spacing = Pt(24)
nf = normal.font
nf.name = "Times New Roman"
nf.size = Pt(12)

# ── Title ─────────────────────────────────────────────────────────────────────
p = add_para(doc, "修改说明", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=120)
p = add_para(doc,
    "Parental Education Level and Neural Development in Young Children: "
    "A Systematic Review",
    italic=True, size=11,
    align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=240)

# Greeting
add_mixed(doc, [("陶老师，您好。", False, False)], before=0, after=80)
add_mixed(doc, [(
    "根据您三轮反馈，我们对文章进行了系统性重写，核心改变是从\u201c文献陈列\u201d转向"
    "\u201c回答领域未解问题\u201d。以下逐条说明。",
    False, False)], before=0, after=200)

# ═══════════════════════════════════════════════════════════════════════════════
# Section 0 – Noble/Farah alignment (新增)
# ═══════════════════════════════════════════════════════════════════════════════
add_para(doc, "零、本文在 Hackman/Noble 未竟方向上的推进", bold=True, size=12,
         before=120, after=80)

add_mixed(doc, [
    ("您三轮反馈的核心锚点是：看看 Hackman & Farah 已有的工作，在前人基础上往前走一步。"
     "我们核查后确认，这一逻辑链已贯穿修改后的全文，现在修改说明中也予以显式标注。",
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ("Hackman & Farah（2009, ", False, False),
    ("Neuroscience & Biobehavioral Reviews", False, True),
    ("）明确呼吁：系统检验 SES 对 0–8 岁儿童神经系统的影响，并分离父母教育与收入的独立贡献。"
     "Noble et al.（2015, ", False, False),
    ("Nature Neuroscience", False, True),
    ("）发现 SES 与皮层面积呈非线性关联，但样本为学龄儿童（6–18 岁），且"
     "未分离父母教育与家庭收入。",
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ("然而，在我们最终检索的 109 篇候选文献中，有 83 篇（76%）因主暴露并非父母教育而被排除"
     "（E2 暴露不符）——",
     False, False),
    ("这是领域 15 年来从未系统响应 Hackman 呼吁的直接文献证据。", True, False),
], before=0, after=60)

add_mixed(doc, [
    ("本综述是首次专门针对父母教育（而非复合 SES）、聚焦 0–8 岁神经发育、整合全部六种神经模态的系统综述。"
     "与 Hackman/Noble 的未竟方向相比，本文的三步推进是：",
     False, False),
], before=0, after=40)

add_mixed(doc, [
    ("（1）", True, False),
    ("跨模态收敛：", True, False),
    ("确认弓状束 FA 是唯一跨研究、跨年龄、控制收入后仍稳健的结构性指标——Noble 2015 未能检验的层面；",
     False, False),
], before=0, after=20, indent_cm=0.8)

add_mixed(doc, [
    ("（2）", True, False),
    ("时序下探：", True, False),
    ("效应在出生后数天即可测（Ramphal 2020），远早于 Noble 2015 的 6 岁起点；",
     False, False),
], before=0, after=20, indent_cm=0.8)

add_mixed(doc, [
    ("（3）", True, False),
    ("系统性盲点揭示：", True, False),
    ("0 篇文献分别检验父亲与母亲教育的独立效应——这是 Hackman 呼吁\"分离父母教育\"15 年后"
     "仍未被响应的领域空白，也是大论文设计的直接依据。",
     False, False),
], before=0, after=120, indent_cm=0.8)

# ═══════════════════════════════════════════════════════════════════════════════
# Section 1
# ═══════════════════════════════════════════════════════════════════════════════
add_para(doc, "一、三条核心知识的改动", bold=True, size=12,
         before=120, after=80)

# ── K1 ──
add_para(doc, "知识一（跨研究收敛性与效应稳健性）", bold=True, italic=False,
         size=12, before=80, after=40)

add_mixed(doc, [(
    "根据您的反馈，我们重新审视了原版知识一的表述：原版使用 "
    "sensitivity/specificity 未能准确反映证据现状——该术语预设了存在诊断准确性研究"
    "（ROC/AUC），但 16 篇文献均为群体水平关联设计，无一项采用预测效度框架。",
    False, False)], before=0, after=60)

add_mixed(doc, [
    ("修改后，知识一回答的是：", False, False),
    ("16 篇中哪个指标跨研究最稳健？", True, False),
    (" 答案是弓状束 FA（", False, False),
    ("arcuate fasciculus fractional anisotropy", False, True),
    ("）：由两个独立团队发现，效应从 8.6 个月（", False, False),
    ("r", False, True),
    (" = 0.48）延伸至 5–8 岁（", False, False),
    ("r", False, True),
    (" = 0.33），且在 5–8 岁样本中控制家庭收入后仍显著——说明教育效应独立于经济资源。",
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ("同时，知识一明确指出：", False, False),
    ("弓状束 FA 的个体预测效度（能否用于筛查高风险儿童？灵敏度/特异度如何？）"
     "从未被检验", True, False),
    ("——这不是本综述的局限，而是领域的空白，本身就是结论。", False, False),
], before=0, after=120)

# ── K2 ──
add_para(doc, "知识二（弓状束是最一致的神经靶点）", bold=True, size=12,
         before=80, after=40)
add_mixed(doc, [
    ("扩写了弓状束的证据链，明确其与阅读获得/阅读障碍的理论连接"
     "（Catani & Mesulam, 2008; Yeatman et al., 2012）。"
     "这是与博士大论文最直接的接口（见下）。", False, False),
], before=0, after=120)

# ── K3 ──
add_para(doc, "知识三（父母不对称性：一个系统性盲点）", bold=True, size=12,
         before=80, after=40)
add_mixed(doc, [
    ("根据您的意见，产前机制（3 篇）不作为独立知识，降级为 §4.2 发育时序分析的 "
     "Phase 1 背景证据（保留其\u300c约束产前路径存在\u300d的功能），但不设独立结论节。",
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ("新知识三基于 16 篇全样本的系统性观察：7 篇仅测母亲教育，9 篇使用合并指数"
     "（均值或最高值），", False, False),
    ("0 篇分别报告父亲与母亲教育的独立效应", True, False),
    ("。这反映的不是统计假象，而是领域一个从未被质疑过的假设——母亲教育已足够代表父母教育。"
     "父亲教育通过不同机制（游戏互动、经济缓冲、语言示范）可能产生独立影响，目前完全未知。",
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ("此空白直接支撑大论文的父母教育四分类设计。", True, False),
], before=0, after=160)

# ═══════════════════════════════════════════════════════════════════════════════
# Section 2 – Three-line table
# ═══════════════════════════════════════════════════════════════════════════════
add_para(doc, "二、与博士大论文的四条接口", bold=True, size=12,
         before=120, after=100)

# Table title
add_mixed(doc, [
    ("表1  本综述为博士大论文提供的理论与方法学支撑", False, False),
], align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=60)

headers = ["接口", "本综述提供", "大论文应用"]
rows = [
    [
        "①弓状束–阅读回路",
        "弓状束 FA 在婴儿期对父母教育敏感，5–8 岁控收入后仍显著",
        "5–6 岁 EEG 的 alpha/theta 振荡是该回路的功能对应，本文提供结构先验依据",
    ],
    [
        "②父母区分设计",
        "0 篇研究分别检验父亲/母亲效应——领域结构性空白",
        "父母教育四分类可分别分析母亲与父亲效应，直接填补此空白",
    ],
    [
        "③EEG 指标先验",
        "Brito & Noble (2020)、Shephard (2019) 等显示 alpha/theta 功率与父母教育关联",
        "选 alpha/theta 作核心指标有文献基础",
    ],
    [
        "④非线性探索基础",
        "16 篇均假设线性；Noble (2015) 非线性发现限于收入–皮层面积，在父母教育–神经的 0–8 岁范围从未验证",
        "四分类设计可初步观察是否存在临界点（如本科是否为阈值），为后续专门检验提供设计参考",
    ],
]

col_widths = [Cm(3.2), Cm(6.5), Cm(6.5)]
tbl = doc.add_table(rows=1 + len(rows), cols=3)
tbl.style = "Table Grid"

# Remove all borders first at table level
_set_table_border(tbl, top_sz="12", bottom_sz="12")

# Header row
hrow = tbl.rows[0]
for j, (cell, hdr) in enumerate(zip(hrow.cells, headers)):
    cell.width = col_widths[j]
    cell_para(cell, [(hdr, True, False)],
              align=WD_ALIGN_PARAGRAPH.CENTER, before=60, after=60)
    # header bottom = thin line
    _set_cell_border(cell, top="none", bottom="single", sz_bottom="6")

# Data rows
for i, row_data in enumerate(rows):
    drow = tbl.rows[i + 1]
    is_last = (i == len(rows) - 1)
    for j, (cell, txt) in enumerate(zip(drow.cells, row_data)):
        cell.width = col_widths[j]
        cell_para(cell, [(txt, False, False)],
                  align=WD_ALIGN_PARAGRAPH.LEFT, before=40, after=40)
        if is_last:
            _set_cell_border(cell, bottom="single", sz_bottom="12")
        else:
            _set_cell_border(cell)   # no borders

# Note under table
p_note = add_para(doc, "", before=60, after=0)
r1 = p_note.add_run("注.")
set_font(r1, size=10, bold=True, italic=True)
r2 = p_note.add_run(
    " alpha/theta = 4–12 Hz EEG 振荡，与丘脑皮层成熟度及语言加工相关；"
    "FA = fractional anisotropy（弓状束各向异性分数）；"
    "四分类 = 高中及以下 / 大专 / 本科 / 研究生。"
)
set_font(r2, size=10)
set_spacing(p_note, before=60, after=120, line=240, rule="auto")

# ═══════════════════════════════════════════════════════════════════════════════
# Section 3
# ═══════════════════════════════════════════════════════════════════════════════
add_para(doc, "三、关于大型队列遗漏的说明", bold=True, size=12,
         before=120, after=80)
add_mixed(doc, [
    ("经逐篇核查：HBCD、GUSTO、HCP-D、Babybrain 等大型队列在检索截止日（2026 年 4 月）"
     "均无已发表的、符合 PICOS 标准的结果论文——仅有研究设计/方案文章，或 SES 测量无法"
     "分离教育效应。16 篇基础完整，无遗漏。", False, False),
], before=0, after=160)

# ═══════════════════════════════════════════════════════════════════════════════
# Summary box (shaded paragraph)
# ═══════════════════════════════════════════════════════════════════════════════
add_para(doc, "核心修改一句话摘要", bold=True, size=12,
         before=120, after=60)

# Shaded paragraph via paragraph border / shading
p_sum = doc.add_paragraph()
set_spacing(p_sum, before=80, after=80, line=360, rule="auto")
p_sum.paragraph_format.left_indent  = Cm(0.5)
p_sum.paragraph_format.right_indent = Cm(0.5)

# Add shading
pPr2 = p_sum._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"),   "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"),  "F2F2F2")
pPr2.append(shd)

run_sum = p_sum.add_run(
    "本文在 Hackman/Noble 未竟方向基础上往前走了三步：（1）弓状束 FA 是 16 篇中唯一"
    "跨研究、跨年龄、控收入后仍稳健的指标，但其个体预测效度从未被检验——这是领域的空白；"
    "（2）该效应 8.6 个月即可测，延伸至 5–8 岁，是阅读障碍研究的神经前体证据；"
    "（3）0 篇研究分别检验父亲 vs 母亲教育效应——此系统性盲点（Hackman 2009 呼吁 15 年"
    "未响应）直接支撑大论文父母教育四分类设计。"
)
set_font(run_sum, size=11)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(out_path)
print(f"Saved: {out_path}")
