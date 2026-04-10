"""
generate_docx_paper02.py
Convert paper02_draft_v1.md → paper02_draft_v1.docx
DCN formatting: 12pt Times New Roman, double-spaced, 2.54cm margins
Tables formatted; inline bold/italic preserved; figures referenced as placeholders
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir  = os.path.dirname(os.path.abspath(__file__))
md_path     = os.path.join(script_dir, "文档_2_paper02_draft_v1.md")
docx_path   = os.path.join(script_dir, "paper02_draft_v1.docx")
figures_dir = os.path.join(script_dir, "figures")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size_pt=12):
    run.bold   = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)

def set_para_spacing(para, space_before=0, space_after=6, line_spacing=2.0):
    """Double-spaced by default, 6pt after paragraph."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(space_before * 20)))
    spacing.set(qn("w:after"),  str(int(space_after  * 20)))
    spacing.set(qn("w:line"),   str(int(line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

def add_para(doc, text, style="Normal", bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False):
    para = doc.add_paragraph(style=style)
    para.alignment = align
    set_para_spacing(para)
    if first_line_indent:
        from docx.shared import Cm
        para.paragraph_format.first_line_indent = Cm(1.27)
    add_inline(para, text, bold=bold, italic=italic)
    return para

def add_inline(para, text, bold=False, italic=False):
    """Parse **bold**, *italic*, ***bold-italic*** inline markup and add runs."""
    # Pattern: ***text***, **text**, *text*, normal text
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|[^*]+)', re.DOTALL)
    for m in pattern.finditer(text):
        chunk = m.group(0)
        if chunk.startswith("***") and chunk.endswith("***"):
            run = para.add_run(chunk[3:-3])
            set_run_font(run, bold=True, italic=True)
        elif chunk.startswith("**") and chunk.endswith("**"):
            run = para.add_run(chunk[2:-2])
            set_run_font(run, bold=True)
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            run = para.add_run(chunk[1:-1])
            set_run_font(run, italic=True)
        else:
            run = para.add_run(chunk)
            set_run_font(run, bold=bold, italic=italic)

def parse_table_md(lines, start):
    """Parse a markdown table starting at lines[start]. Returns (table_data, next_idx)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if not all(re.match(r"^[-: ]+$", c) for c in cells):
                rows.append(cells)
        elif rows:
            break
        i += 1
    return rows, i

def add_table(doc, rows):
    """Add a Word table from list-of-lists rows (first row = header)."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # pad rows
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            # Clear default paragraph and add new one
            para = cell.paragraphs[0]
            para.clear()
            is_header = ri == 0
            add_inline(para, cell_text, bold=is_header)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            from docx.oxml.ns import qn
            pPr = para._p.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"),  "0")
            spacing.set(qn("w:line"),   str(int(1.15 * 240)))
            spacing.set(qn("w:lineRule"), "auto")
            pPr.append(spacing)
    doc.add_paragraph()  # space after table

# ── Document setup ─────────────────────────────────────────────────────────

doc = Document()

# Page margins: 2.54 cm all sides
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# Default Normal style
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)

# Heading styles
for level, (name, size, bold) in enumerate([
    ("Heading 1", 14, True),
    ("Heading 2", 13, True),
    ("Heading 3", 12, True),
    ("Heading 4", 12, True),
], start=1):
    try:
        s = doc.styles[f"Heading {level}"]
    except KeyError:
        s = doc.styles.add_style(f"Heading {level}", WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = "Times New Roman"
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)

# ── Parse markdown ────────────────────────────────────────────────────────────

with open(md_path, encoding="utf-8") as f:
    raw = f.read()

lines = raw.split("\n")

skip_meta_lines = 6  # skip the draft-version/status/target-journal header block
i = 0
in_blockquote = False
blockquote_lines = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip empty
    if not stripped:
        i += 1
        continue

    # Skip horizontal rules
    if stripped in ("---", "***", "___"):
        i += 1
        continue

    # Skip metadata lines (lines 1-6: draft version etc.)
    if i < skip_meta_lines and (stripped.startswith("**Draft") or
                                 stripped.startswith("**Status") or
                                 stripped.startswith("**Target") or
                                 stripped == "---"):
        i += 1
        continue

    # Skip the "(To be compiled...)" placeholder in References
    if stripped.startswith("*(To be compiled") or stripped.startswith("*(Methods are largely"):
        i += 1
        continue

    # H1 title
    if stripped.startswith("# ") and not stripped.startswith("## "):
        title_text = stripped[2:]
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(para, space_before=0, space_after=12, line_spacing=1.5)
        run = para.add_run(title_text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        i += 1
        continue

    # H2
    if stripped.startswith("## ") and not stripped.startswith("### "):
        heading = stripped[3:]
        para = doc.add_heading(heading, level=1)
        set_para_spacing(para, space_before=12, space_after=3, line_spacing=1.5)
        for run in para.runs:
            run.font.name = "Times New Roman"; run.font.size = Pt(14); run.font.bold = True
        i += 1
        continue

    # H3
    if stripped.startswith("### ") and not stripped.startswith("#### "):
        heading = stripped[4:]
        para = doc.add_heading(heading, level=2)
        set_para_spacing(para, space_before=10, space_after=2, line_spacing=1.5)
        for run in para.runs:
            run.font.name = "Times New Roman"; run.font.size = Pt(13); run.font.bold = True
        i += 1
        continue

    # H4
    if stripped.startswith("#### "):
        heading = stripped[5:]
        para = doc.add_heading(heading, level=3)
        set_para_spacing(para, space_before=8, space_after=2, line_spacing=1.5)
        for run in para.runs:
            run.font.name = "Times New Roman"; run.font.size = Pt(12); run.font.bold = True
        i += 1
        continue

    # Blockquote (search strategy)
    if stripped.startswith("> "):
        if not in_blockquote:
            in_blockquote = True
            blockquote_lines = []
        blockquote_lines.append(stripped[2:])
        # check if next line is also blockquote
        if i + 1 < len(lines) and lines[i+1].strip().startswith("> "):
            i += 1
            continue
        else:
            # flush blockquote
            bq_text = " ".join(blockquote_lines)
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.left_indent  = Cm(1.27)
            para.paragraph_format.right_indent = Cm(1.27)
            set_para_spacing(para, space_before=3, space_after=3, line_spacing=1.15)
            add_inline(para, bq_text, italic=True)
            in_blockquote = False
            blockquote_lines = []
            i += 1
            continue

    # Table
    if stripped.startswith("|"):
        rows, next_i = parse_table_md(lines, i)
        if rows:
            add_table(doc, rows)
            i = next_i
            continue

    # Otherwise normal paragraph
    add_para(doc, stripped, first_line_indent=False)
    i += 1

# ── Save ───────────────────────────────────────────────────────────────────

doc.save(docx_path)
print(f"Saved: {docx_path}")
