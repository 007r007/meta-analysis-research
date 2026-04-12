# -*- coding: utf-8 -*-
"""
gen_rev_note.py
Generate revision_note_to_supervisor.docx
Three-line table, Times New Roman 12pt, double spacing
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# output next to this script (no Chinese in path)
script_dir = os.path.dirname(os.path.abspath(__file__))
out_path   = os.path.join(
    script_dir,
    '\u8001\u5e08\u6307\u5bfc\u540e\u5b8c\u5584',   # 老师指导后完善
    'revision_note_to_supervisor.docx'
)

# ── helpers ────────────────────────────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic


def clear_spacing(pPr):
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)


def set_spacing(para, before=0, after=0, line=480, rule='auto'):
    pPr = para._p.get_or_add_pPr()
    clear_spacing(pPr)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'),   str(before))
    sp.set(qn('w:after'),    str(after))
    sp.set(qn('w:line'),     str(line))
    sp.set(qn('w:lineRule'), rule)
    pPr.append(sp)


def add_para(doc, text='', bold=False, italic=False, size=12,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0,
             line=480, rule='auto', indent_cm=0):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after, line=line, rule=rule)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_mixed(doc, parts, before=0, after=0, line=480,
              align=WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after, line=line)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


# ── Three-line table helpers ───────────────────────────────────────────────────

def _set_table_border(tbl, top_sz='12', bottom_sz='12'):
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        if side == 'top':
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    top_sz)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
        elif side == 'bottom':
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    bottom_sz)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
        else:
            b.set(qn('w:val'),   'none')
            b.set(qn('w:sz'),    '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_cell_border(cell, bottom='none', sz_bottom='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        if side == 'bottom' and bottom != 'none':
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    sz_bottom)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
        else:
            b.set(qn('w:val'),   'none')
            b.set(qn('w:sz'),    '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def cell_para(cell, parts, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT,
              before=40, after=40):
    from docx.text.paragraph import Paragraph
    # clear existing content
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    tc = cell._tc
    p_elem = OxmlElement('w:p')
    tc.append(p_elem)
    p = Paragraph(p_elem, cell)
    p.alignment = align
    set_spacing(p, before=before, after=after, line=240, rule='auto')
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.18)
    section.right_margin  = Cm(3.18)

normal = doc.styles['Normal']
normal.paragraph_format.line_spacing = Pt(24)
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

# ── Title ──────────────────────────────────────────────────────────────────────
add_para(doc, '\u4fee\u6539\u8bf4\u660e',   # 修改说明
         bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
         before=0, after=120)
add_para(doc,
         'Parental Education Level and Neural Development in Young Children: '
         'A Systematic Review',
         italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
         before=0, after=240)

# Greeting
add_para(doc, '\u9676\u8001\u5e08\uff0c\u60a8\u597d\u3002',   # 陶老师，您好。
         before=0, after=80)

# Note: Chinese curly quotes replaced with ASCII equivalents in Python strings
add_para(doc,
         '\u6839\u636e\u60a8\u4e09\u8f6e\u53cd\u9988\uff0c\u6211\u4eec\u5bf9\u6587\u7ae0'
         '\u8fdb\u884c\u4e86\u7cfb\u7edf\u6027\u91cd\u5199\uff0c\u6838\u5fc3\u6539\u53d8'
         '\u662f\u4ece\u201c\u6587\u732e\u9648\u5217\u201d\u8f6c\u5411\u201c\u56de\u7b54'
         '\u9886\u57df\u672a\u89e3\u95ee\u9898\u201d\u3002\u4ee5\u4e0b\u9010\u6761\u8bf4\u660e\u3002',
         before=0, after=200)

# ── Section 1 ──────────────────────────────────────────────────────────────────
add_para(doc, '\u4e00\u3001\u4e09\u6761\u6838\u5fc3\u77e5\u8bc6\u7684\u6539\u52a8',  # 一、三条核心知识的改动
         bold=True, size=12, before=120, after=80)

# K1 heading
add_para(doc,
         '\u77e5\u8bc6\u4e00\uff08\u8de8\u7814\u7a76\u6536\u655b\u6027\u4e0e\u6548\u5e94'
         '\u7a33\u5065\u6027\uff09',  # 知识一（跨研究收敛性与效应稳健性）
         bold=True, size=12, before=80, after=40)

add_para(doc,
         '\u6839\u636e\u60a8\u7684\u53cd\u9988\uff0c\u6211\u4eec\u91cd\u65b0\u5ba1\u89c6'
         '\u4e86\u539f\u7248\u77e5\u8bc6\u4e00\u7684\u8868\u8ff0\uff1a\u539f\u7248\u4f7f'
         '\u7528 sensitivity/specificity \u672a\u80fd\u51c6\u786e\u53cd\u6620\u8bc1\u636e'
         '\u73b0\u72b0\u2014\u2014\u8be5\u672f\u8bed\u9884\u8bbe\u4e86\u5b58\u5728\u8bca'
         '\u65ad\u51c6\u786e\u6027\u7814\u7a76\uff08ROC/AUC\uff09\uff0c\u4f46 16 \u7bc7'
         '\u6587\u732e\u5747\u4e3a\u7fa4\u4f53\u6c34\u5e73\u5173\u8054\u8bbe\u8ba1\uff0c'
         '\u65e0\u4e00\u9879\u91c7\u7528\u9884\u6d4b\u6548\u5ea6\u6846\u67b6\u3002',
         before=0, after=60)

add_mixed(doc, [
    ('\u4fee\u6539\u540e\uff0c\u77e5\u8bc6\u4e00\u56de\u7b54\u7684\u662f\uff1a', False, False),  # 修改后，知识一回答的是：
    ('16 \u7bc7\u4e2d\u54ea\u4e2a\u6307\u6807\u8de8\u7814\u7a76\u6700\u7a33\u5065\uff1f',  # 16篇中哪个指标跨研究最稳健？
     True, False),
    (' \u7b54\u6848\u662f\u5f13\u72b6\u675f FA\uff08', False, False),  # 答案是弓状束 FA（
    ('arcuate fasciculus fractional anisotropy', False, True),
    ('\uff09\uff1a\u7531\u4e24\u4e2a\u72ec\u7acb\u56e2\u961f\u53d1\u73b0\uff0c\u6548\u5e94\u4ece 8.6 \u4e2a\u6708\uff08',
     False, False),  # ）：由两个独立团队发现，效应从 8.6 个月（
    ('r', False, True),
    (' = 0.48\uff09\u5ef6\u4f38\u81f3 5\u20138 \u5c81\uff08', False, False),  # = 0.48）延伸至 5–8 岁（
    ('r', False, True),
    (' = 0.33\uff09\uff0c\u4e14\u5728 5\u20138 \u5c81\u6837\u672c\u4e2d\u63a7\u5236\u5bb6\u5ead\u6536\u5165\u540e'
     '\u4ecd\u663e\u8457\u2014\u2014\u8bf4\u660e\u6559\u80b2\u6548\u5e94\u72ec\u7acb\u4e8e\u7ecf\u6d4e\u8d44\u6e90\u3002',
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ('\u540c\u65f6\uff0c\u77e5\u8bc6\u4e00\u660e\u786e\u6307\u51fa\uff1a', False, False),  # 同时，知识一明确指出：
    ('\u5f13\u72b6\u675f FA \u7684\u4e2a\u4f53\u9884\u6d4b\u6548\u5ea6\uff08\u80fd\u5426\u7528\u4e8e\u7b5b\u67e5'
     '\u9ad8\u98ce\u9669\u513f\u7ae5\uff1f\u7075\u654f\u5ea6/\u7279\u5f02\u5ea6\u5982\u4f55\uff1f\uff09'
     '\u4ece\u672a\u88ab\u68c0\u9a8c',
     True, False),
    ('\u2014\u2014\u8fd9\u4e0d\u662f\u672c\u7efc\u8ff0\u7684\u5c40\u9650\uff0c\u800c\u662f\u9886\u57df\u7684'
     '\u7a7a\u767d\uff0c\u672c\u8eab\u5c31\u662f\u7ed3\u8bba\u3002',
     False, False),
], before=0, after=120)

# K2 heading
add_para(doc,
         '\u77e5\u8bc6\u4e8c\uff08\u5f13\u72b6\u675f\u662f\u6700\u4e00\u81f4\u7684\u795e\u7ecf\u9776\u70b9\uff09',
         bold=True, size=12, before=80, after=40)
add_para(doc,
         '\u6269\u5199\u4e86\u5f13\u72b6\u675f\u7684\u8bc1\u636e\u94fe\uff0c\u660e\u786e\u5176\u4e0e\u9605\u8bfb'
         '\u83b7\u5f97/\u9605\u8bfb\u969c\u788d\u7684\u7406\u8bba\u8fde\u63a5\uff08Catani & Mesulam, 2008; '
         'Yeatman et al., 2012\uff09\u3002\u8fd9\u662f\u4e0e\u535a\u58eb\u5927\u8bba\u6587\u6700\u76f4\u63a5'
         '\u7684\u63a5\u53e3\uff08\u89c1\u4e0b\uff09\u3002',
         before=0, after=120)

# K3 heading
add_para(doc,
         '\u77e5\u8bc6\u4e09\uff08\u7236\u6bcd\u4e0d\u5bf9\u79f0\u6027\uff1a\u4e00\u4e2a\u7cfb\u7edf\u6027\u76f2\u70b9\uff09',
         bold=True, size=12, before=80, after=40)
add_para(doc,
         '\u6839\u636e\u60a8\u7684\u610f\u89c1\uff0c\u4ea7\u524d\u673a\u5236\uff083 \u7bc7\uff09\u4e0d\u4f5c\u4e3a'
         '\u72ec\u7acb\u77e5\u8bc6\uff0c\u964d\u7ea7\u4e3a \u00a74.2 \u53d1\u80b2\u65f6\u5e8f\u5206\u6790\u7684'
         ' Phase 1 \u80cc\u666f\u8bc1\u636e\uff08\u4fdd\u7559\u5176\u201c\u7ea6\u675f\u4ea7\u524d\u8def\u5f84'
         '\u5b58\u5728\u201d\u7684\u529f\u80fd\uff09\uff0c\u4f46\u4e0d\u8bbe\u72ec\u7acb\u7ed3\u8bba\u8282\u3002',
         before=0, after=60)

add_mixed(doc, [
    ('\u65b0\u77e5\u8bc6\u4e09\u57fa\u4e8e 16 \u7bc7\u5168\u6837\u672c\u7684\u7cfb\u7edf\u6027\u89c2\u5bdf\uff1a'
     '7 \u7bc7\u4ec5\u6d4b\u6bcd\u4eb2\u6559\u80b2\uff0c9 \u7bc7\u4f7f\u7528\u5408\u5e76\u6307\u6570'
     '\uff08\u5747\u503c\u6216\u6700\u9ad8\u503c\uff09\uff0c',
     False, False),
    ('0 \u7bc7\u5206\u522b\u62a5\u544a\u7236\u4eb2\u4e0e\u6bcd\u4eb2\u6559\u80b2\u7684\u72ec\u7acb\u6548\u5e94',
     True, False),
    ('\u3002\u8fd9\u53cd\u6620\u7684\u4e0d\u662f\u7edf\u8ba1\u5047\u8c61\uff0c\u800c\u662f\u9886\u57df\u4e00\u4e2a'
     '\u4ece\u672a\u88ab\u8d28\u7591\u8fc7\u7684\u5047\u8bbe\u2014\u2014\u6bcd\u4eb2\u6559\u80b2\u5df2\u8db3\u591f'
     '\u4ee3\u8868\u7236\u6bcd\u6559\u80b2\u3002\u7236\u4eb2\u6559\u80b2\u901a\u8fc7\u4e0d\u540c\u673a\u5236'
     '\uff08\u6e38\u620f\u4e92\u52a8\u3001\u7ecf\u6d4e\u7f13\u51b2\u3001\u8bed\u8a00\u793a\u8303\uff09\u53ef\u80fd'
     '\u4ea7\u751f\u72ec\u7acb\u5f71\u54cd\uff0c\u76ee\u524d\u5b8c\u5168\u672a\u77e5\u3002',
     False, False),
], before=0, after=60)

add_mixed(doc, [
    ('\u6b64\u7a7a\u767d\u76f4\u63a5\u652f\u6491\u5927\u8bba\u6587\u7684\u7236\u6bcd\u6559\u80b2\u56db\u5206\u7c7b\u8bbe\u8ba1\u3002',
     True, False),
], before=0, after=160)

# ── Section 2: Three-line table ────────────────────────────────────────────────
add_para(doc,
         '\u4e8c\u3001\u4e0e\u535a\u58eb\u5927\u8bba\u6587\u7684\u56db\u6761\u63a5\u53e3',
         bold=True, size=12, before=120, after=100)

add_para(doc,
         '\u8868 1  \u672c\u7efc\u8ff0\u4e3a\u535a\u58eb\u5927\u8bba\u6587\u63d0\u4f9b\u7684\u7406\u8bba\u4e0e\u65b9\u6cd5\u5b66\u652f\u6491',
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=60)

headers = [
    '\u63a5\u53e3',   # 接口
    '\u672c\u7efc\u8ff0\u63d0\u4f9b',   # 本综述提供
    '\u5927\u8bba\u6587\u5e94\u7528',   # 大论文应用
]
rows = [
    [
        '\u2460\u5f13\u72b6\u675f\u2013\u9605\u8bfb\u56de\u8def',   # ①弓状束–阅读回路
        '\u5f13\u72b6\u675f FA \u5728\u5a74\u513f\u671f\u5bf9\u7236\u6bcd\u6559\u80b2\u654f\u611f\uff0c5\u20138 \u5c81\u63a7\u6536\u5165\u540e\u4ecd\u663e\u8457',
        '5\u20136 \u5c81 EEG \u7684 alpha/theta \u632f\u8361\u662f\u8be5\u56de\u8def\u7684\u529f\u80fd\u5bf9\u5e94\uff0c\u672c\u6587\u63d0\u4f9b\u7ed3\u6784\u5148\u9a8c\u4f9d\u636e',
    ],
    [
        '\u2461\u7236\u6bcd\u533a\u5206\u8bbe\u8ba1',   # ②父母区分设计
        '0 \u7bc7\u7814\u7a76\u5206\u522b\u68c0\u9a8c\u7236\u4eb2/\u6bcd\u4eb2\u6548\u5e94\u2014\u2014\u9886\u57df\u7ed3\u6784\u6027\u7a7a\u767d',
        '\u7236\u6bcd\u6559\u80b2\u56db\u5206\u7c7b\u53ef\u5206\u522b\u5206\u6790\u6bcd\u4eb2\u4e0e\u7236\u4eb2\u6548\u5e94\uff0c\u76f4\u63a5\u586b\u8865\u6b64\u7a7a\u767d',
    ],
    [
        '\u2462EEG \u6307\u6807\u5148\u9a8c',   # ③EEG 指标先验
        'Brito & Noble (2020)\u3001Shephard (2019) \u7b49\u663e\u793a alpha/theta \u529f\u7387\u4e0e\u7236\u6bcd\u6559\u80b2\u5173\u8054',
        '\u9009 alpha/theta \u4f5c\u6838\u5fc3\u6307\u6807\u6709\u6587\u732e\u57fa\u7840',
    ],
    [
        '\u2463\u975e\u7ebf\u6027\u63a2\u7d22\u57fa\u7840',   # ④非线性探索基础
        '16 \u7bc7\u5747\u5047\u8bbe\u7ebf\u6027\uff1bNoble (2015) \u975e\u7ebf\u6027\u53d1\u73b0\u9650\u4e8e\u6536\u5165\u2013\u76ae\u5c42\u9762\u79ef\uff0c\u5728\u7236\u6bcd\u6559\u80b2\u2013\u795e\u7ecf\u7684 0\u20138 \u5c81\u8303\u56f4\u4ece\u672a\u9a8c\u8bc1',
        '\u56db\u5206\u7c7b\u8bbe\u8ba1\u53ef\u521d\u6b65\u89c2\u5bdf\u662f\u5426\u5b58\u5728\u4e34\u754c\u70b9\uff0c\u4e3a\u540e\u7eed\u4e13\u95e8\u68c0\u9a8c\u63d0\u4f9b\u8bbe\u8ba1\u53c2\u8003',
    ],
]

col_widths = [Cm(3.0), Cm(6.5), Cm(6.5)]
tbl = doc.add_table(rows=1 + len(rows), cols=3)
tbl.style = 'Table Grid'
_set_table_border(tbl, top_sz='12', bottom_sz='12')

# Header row
hrow = tbl.rows[0]
for j, (cell, hdr) in enumerate(zip(hrow.cells, headers)):
    cell.width = col_widths[j]
    cell_para(cell, [(hdr, True, False)],
              align=WD_ALIGN_PARAGRAPH.CENTER, before=60, after=60)
    _set_cell_border(cell, bottom='single', sz_bottom='6')

# Data rows
for i, row_data in enumerate(rows):
    drow = tbl.rows[i + 1]
    is_last = (i == len(rows) - 1)
    for j, (cell, txt) in enumerate(zip(drow.cells, row_data)):
        cell.width = col_widths[j]
        cell_para(cell, [(txt, False, False)],
                  align=WD_ALIGN_PARAGRAPH.LEFT, before=40, after=40)
        if is_last:
            _set_cell_border(cell, bottom='single', sz_bottom='12')
        else:
            _set_cell_border(cell)

# Table note
p_note = add_para(doc, '', before=60, after=120, line=240, rule='auto')
r1 = p_note.add_run('\u6ce8.')   # 注.
set_font(r1, size=10, bold=True, italic=True)
r2 = p_note.add_run(
    ' alpha/theta = 4\u201312 Hz EEG \u632f\u8361\uff0c\u4e0e\u4e18\u8111\u76ae\u5c42'
    '\u6210\u719f\u5ea6\u53ca\u8bed\u8a00\u52a0\u5de5\u76f8\u5173\uff1b'
    'FA = fractional anisotropy\uff08\u5f13\u72b6\u675f\u5404\u5411\u5f02\u6027\u5206\u6570\uff09\uff1b'
    '\u56db\u5206\u7c7b = \u9ad8\u4e2d\u53ca\u4ee5\u4e0b / \u5927\u4e13 / \u672c\u79d1 / \u7814\u7a76\u751f\u3002'
)
set_font(r2, size=10)

# ── Section 3 ──────────────────────────────────────────────────────────────────
add_para(doc,
         '\u4e09\u3001\u5173\u4e8e\u5927\u578b\u961f\u5217\u9057\u6f0f\u7684\u8bf4\u660e',
         bold=True, size=12, before=120, after=80)
add_para(doc,
         '\u7ecf\u9010\u7bc7\u6838\u67e5\uff1aHBCD\u3001GUSTO\u3001HCP-D\u3001Babybrain '
         '\u7b49\u5927\u578b\u961f\u5217\u5728\u68c0\u7d22\u622a\u6b62\u65e5\uff082026 \u5e74 4 \u6708\uff09'
         '\u5747\u65e0\u5df2\u53d1\u8868\u7684\u3001\u7b26\u5408 PICOS \u6807\u51c6\u7684\u7ed3\u679c\u6027\u8bba\u6587'
         '\u2014\u2014\u4ec5\u6709\u7814\u7a76\u8bbe\u8ba1/\u65b9\u6848\u6587\u7ae0\uff0c\u6216 SES '
         '\u6d4b\u91cf\u65e0\u6cd5\u5206\u79bb\u6559\u80b2\u6548\u5e94\u3002'
         '16 \u7bc7\u57fa\u7840\u5b8c\u6574\uff0c\u65e0\u9057\u6f0f\u3002',
         before=0, after=160)

# ── Summary box ────────────────────────────────────────────────────────────────
add_para(doc,
         '\u6838\u5fc3\u4fee\u6539\u4e00\u53e5\u8bdd\u6458\u8981',
         bold=True, size=12, before=120, after=60)

p_sum = doc.add_paragraph()
set_spacing(p_sum, before=80, after=80, line=360, rule='auto')
p_sum.paragraph_format.left_indent  = Cm(0.5)
p_sum.paragraph_format.right_indent = Cm(0.5)
pPr2 = p_sum._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'),   'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'),  'F2F2F2')
pPr2.append(shd)
run_sum = p_sum.add_run(
    '\u4fee\u6539\u540e\u4e09\u6761\u77e5\u8bc6\uff1a'
    '\uff081\uff09\u5f13\u72b6\u675f FA \u662f 16 \u7bc7\u4e2d\u552f\u4e00\u8de8\u7814\u7a76\u3001'
    '\u8de8\u5e74\u9f84\u3001\u63a7\u6536\u5165\u540e\u4ecd\u7a33\u5065\u7684\u6307\u6807\uff0c'
    '\u4f46\u5176\u4e2a\u4f53\u9884\u6d4b\u6548\u5ea6\u4ece\u672a\u88ab\u68c0\u9a8c'
    '\u2014\u2014\u8fd9\u662f\u9886\u57df\u7684\u7a7a\u767d\uff1b'
    '\uff082\uff09\u8be5\u6548\u5e94 8.6 \u4e2a\u6708\u5373\u53ef\u6d4b\uff0c\u5ef6\u4f38\u81f3'
    ' 5\u20138 \u5c81\uff0c\u662f\u9605\u8bfb\u969c\u788d\u7814\u7a76\u7684\u795e\u7ecf\u524d\u4f53\u8bc1\u636e\uff1b'
    '\uff083\uff090 \u7bc7\u7814\u7a76\u5206\u522b\u68c0\u9a8c\u7236\u4eb2 vs \u6bcd\u4eb2\u6559\u80b2'
    '\u6548\u5e94\u2014\u2014\u6b64\u7cfb\u7edf\u6027\u76f2\u70b9\u76f4\u63a5\u652f\u6491\u5927\u8bba\u6587'
    '\u7236\u6bcd\u6559\u80b2\u56db\u5206\u7c7b\u8bbe\u8ba1\u3002'
)
set_font(run_sum, size=11)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(out_path)
print(f'Saved: {out_path}')
