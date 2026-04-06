"""
通用投稿前完整性审阅脚本 v2
Universal Pre-Submission Checklist for Systematic Review Papers
==============================================================
适用范围：叙述性系统综述 / 元分析 论文(.docx格式)
基于 Paper2 审阅框架重构，v2 修复了 v1 的主要检测缺陷

v1 → v2 改进：
  - Part 4: 引用regex同时捕获 (Author, Year) 和 Author (Year) 两种格式
  - Part 4: 参考文献列表提取逻辑加固（不依赖首字母大写）
  - Part 4: 新增 orphan citation 双向检测（正文有但列表无 / 列表有但正文未引）
  - Part 5: 连字符检测排除DOI路径，减少误报
  - Part 5: 统计符号输出前3条语境片段，便于人工确认
  - Part 6: 新增 APA 数字格式检查（前导零、效应量异常值）
  - Part 7: Figure/Table caption 提取改为全文搜索，不再依赖行首锚定
  - Part 7: 新增正文总字数统计（排除摘要和参考文献）
  - Part 7: Table 同样做 caption vs 引用的集合比对

使用方法：
  1. 修改下方 CONFIG 区域中的论文专属参数
  2. 运行: python presubmit_check_universal_v2.py

作者：CC (Claude Code)
创建：2026-04-06，v2 更新：2026-04-06
适用项目：paper-01, paper-02, paper-03...
"""

import sys
import re
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

# ═══════════════════════════════════════════════════════════════════
# ★ CONFIG — 每篇论文修改这里 ★
# ═══════════════════════════════════════════════════════════════════
CONFIG = {
    # 文档路径（绝对路径）
    "doc_path": r"E:\Meta-analysis writing project\projects\paper-01\06-write\paper01_draft_v1.docx",

    # 论文标识（用于报告标题）
    "paper_id": "Paper-01",
    "paper_topic": "WM Training in Healthy Older Adults",

    # PRISMA 关键数字（从筛选Excel核实后填入）
    "n_total_records":  4168,   # 去重后总记录数
    "n_fulltext":        445,   # 进入全文筛选数
    "n_included":         56,   # 最终纳入数
    "n_excluded_stage12": 3723, # 标题摘要阶段排除数

    # 使用的数据库（list）
    "databases": ["PubMed", "PsycINFO", "Web of Science", "Scopus"],

    # 预期表格数量
    "expected_tables": 3,

    # 预期章节（按论文实际结构调整）
    "expected_sections": [
        "Abstract", "Introduction", "Methods",
        "Results", "Discussion", "Conclusion", "References"
    ],

    # 论文主题相关词（用于摘要背景检查）
    "topic_keywords": ["working memory", "older adult", "cognitive training"],

    # 内部代码格式（筛选阶段用的临时标记，投稿前应清除）
    # 填 None 则跳过此检查
    "internal_code_pattern": None,   # e.g. r'\[FT\d{2}\]' or r'\bseq\d+\b'

    # 摘要字数上限（根据目标期刊调整）
    "abstract_word_limit": 300,

    # 正文总字数上限（不含摘要和参考文献）；填 None 则只统计不报警
    "body_word_limit": 10000,

    # 需要核查的关键声明（缺失时报警）
    "required_statements": {
        "PRISMA 2020":              r"PRISMA",
        "Registration/PROSPERO":    r"PROSPERO|pre.regist|not pre.regist",
        "SWiM/Narrative synthesis": r"SWiM|narrative synthesis",
        "RoB tool cited":           r"RoB 2|RoB2|ROBINS.I|Sterne|risk of bias tool",
        "Limitations section":      r"[Ll]imitation",
        "Publication bias":         r"publication bias",
        "Conflict of interest":     r"conflict of interest|competing interest",
        "Funding statement":        r"funding|grant|no funding|no financial",
        "Data availability":        r"data avail|supplementary|supporting information",
        "Ethics/IRB":               r"ethics|IRB|institutional review|ethical approval",
    },

    # 语言过强词（建议软化）
    "overstrong_language": [
        "conclusive",
        "clearly demonstrates",
        "proves",
        "definitively",
        "unambiguously",
    ],
}
# ═══════════════════════════════════════════════════════════════════


# ── helpers ─────────────────────────────────────────────────────────

def load_doc(path):
    doc = Document(path)
    # 主体段落
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    # 表格单元格（含嵌套）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text
    return doc, all_text


def extract_section(doc, start_marker, stop_markers):
    """提取从 start_marker 段落开始、到 stop_markers 之一结束的文本（不含标题行）。"""
    text = ""
    collecting = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == start_marker:
            collecting = True
            continue
        if collecting:
            if any(t == m or t.startswith(m) for m in stop_markers):
                break
            text += para.text + " "
    return text


def extract_references(doc):
    """提取 References 节后所有非空段落，用年份正则过滤掉非引用行。"""
    entries = []
    ref_start = False
    for para in doc.paragraphs:
        if para.text.strip() == "References":
            ref_start = True
            continue
        if ref_start:
            t = para.text.strip()
            if not t:
                continue
            # 保留含年份的段落（排除纯章节标题）
            if re.search(r'\b(19|20)\d{2}\b', t):
                entries.append(t)
    return entries


def extract_author_year_keys(text):
    """从文本中提取所有引用的 (FirstAuthor, Year) 键，兼容两种APA格式：
    括号式: (Smith, 2020) / (Smith et al., 2020) / (Smith & Jones, 2020)
    叙述式: Smith (2020) / Smith et al. (2020)
    """
    keys = set()
    # 括号式: (Author..., Year)
    p1 = re.findall(
        r'\(([A-Z][a-zA-Z\u00c0-\u024f\-]+)(?:\s+et al\.|\s+&\s+[A-Z][a-z]+|\s+and\s+[A-Z][a-z]+)?'
        r'(?:,\s*|\s+)(\d{4}[a-z]?)\)',
        text
    )
    # 叙述式: Author (Year) / Author et al. (Year)
    p2 = re.findall(
        r'\b([A-Z][a-zA-Z\u00c0-\u024f\-]+)(?:\s+et al\.)?\s+\((\d{4}[a-z]?)\)',
        text
    )
    for author, year in p1 + p2:
        keys.add((author.strip(), year.strip()))
    return keys


def extract_ref_list_keys(ref_entries):
    """从参考文献列表条目中提取 (FirstAuthor, Year) 键。
    APA格式: Author, A. A. (Year). Title...
    """
    keys = set()
    for entry in ref_entries:
        # 匹配行首的姓氏（第一个逗号前）和括号年份
        m = re.match(r'^([A-Z][a-zA-Z\u00c0-\u024f\-]+)(?:,|\s)', entry)
        y = re.search(r'\((\d{4}[a-z]?)\)', entry)
        if m and y:
            keys.add((m.group(1), y.group(1)))
    return keys


def context_snippet(text, pattern, n=3, window=50):
    """返回 pattern 匹配位置周围 ±window 字符的片段，最多 n 条。"""
    snippets = []
    for m in re.finditer(pattern, text):
        start = max(0, m.start() - window)
        end   = min(len(text), m.end() + window)
        snippets.append(f"...{text[start:end].strip()}...")
        if len(snippets) >= n:
            break
    return snippets


# ── Part 1 ───────────────────────────────────────────────────────────

def check_part1_completeness(doc, all_text, cfg):
    print("\n[PART 1] Content Completeness")
    print("-" * 65)

    # 1a. Internal codes
    pattern = cfg.get("internal_code_pattern")
    if pattern:
        codes = re.findall(pattern, all_text)
        if codes:
            print(f"  [WARN] Internal codes present ({len(codes)}): {codes[:5]}")
        else:
            print(f"  [OK] No internal codes (pattern: {pattern})")
    else:
        print("  [SKIP] internal_code_pattern not configured")

    # 1b. Generic placeholders
    print("\n  Placeholder check:")
    placeholders = ['XXX', 'TODO', 'TBD', '[INSERT', 'PLACEHOLDER', '???',
                    '[Author names to be added]', 'forthcoming', 'in press']
    found_any = False
    for ph in placeholders:
        if ph.lower() in all_text.lower():
            count = all_text.lower().count(ph.lower())
            print(f"    [WARN] '{ph}': {count} times")
            found_any = True
    if not found_any:
        print("    [OK] No placeholders found")

    # 1c. Section structure
    print("\n  Section structure:")
    for s in cfg["expected_sections"]:
        found = any(p.text.strip() == s or p.text.strip().endswith(s)
                    for p in doc.paragraphs)
        print(f"    {'[OK]' if found else '[MISSING]'} {s}")

    # 1d. Tables
    n_tables = len(doc.tables)
    exp = cfg["expected_tables"]
    status = "[OK]" if n_tables == exp else f"[WARN: expected {exp}]"
    print(f"\n  Tables: {n_tables} found {status}")
    for i, tbl in enumerate(doc.tables):
        print(f"    Table {i+1}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")


# ── Part 2 ───────────────────────────────────────────────────────────

def check_part2_numbers(doc, all_text, cfg):
    print("\n[PART 2] Key Numbers Verification")
    print("-" * 65)

    n_total  = cfg["n_total_records"]
    n_ft     = cfg["n_fulltext"]
    n_inc    = cfg["n_included"]
    n_excl12 = cfg["n_excluded_stage12"]

    # PRISMA arithmetic
    calc_ft      = n_total - n_excl12
    calc_excl_ft = n_ft - n_inc
    print("  PRISMA arithmetic:")
    print(f"    Total({n_total:,}) - ExclStage12({n_excl12:,}) = {calc_ft:,} "
          f"{'== fulltext OK' if calc_ft == n_ft else f'!= fulltext({n_ft}) [MISMATCH]'}")
    print(f"    Fulltext({n_ft}) - Included({n_inc}) = {calc_excl_ft} (= full-text excluded)")

    # Numbers in text — handle comma-formatted numbers (e.g. "4,168")
    print("\n  Key numbers in text:")
    for label, num in [
        (f"Total records ({n_total:,})", n_total),
        (f"Fulltext ({n_ft})",           n_ft),
        (f"Included ({n_inc})",          n_inc),
    ]:
        # Match both "4168" and "4,168"
        s = str(num)
        if len(s) > 3:
            pat = s[:-3] + '[,.]?' + s[-3:]
        else:
            pat = s
        count = len(re.findall(rf'\b{pat}\b', all_text))
        print(f"    {label}: {count} mentions {'[OK]' if count >= 2 else '[CHECK: few mentions]'}")

    # Databases
    print("\n  Databases:")
    for db in cfg["databases"]:
        n = len(re.findall(re.escape(db), all_text))
        print(f"    {db}: {n} times {'[OK]' if n > 0 else '[MISSING]'}")


# ── Part 3 ───────────────────────────────────────────────────────────

def check_part3_abstract(doc, all_text, cfg):
    print("\n[PART 3] Abstract")
    print("-" * 65)

    abstract_text = extract_section(doc, "Abstract", ["1.", "1. Introduction"])
    wc    = len(abstract_text.split())
    limit = cfg["abstract_word_limit"]
    print(f"  Word count: {wc} {'[OK]' if wc <= limit else f'[WARN: exceeds {limit}]'}")

    print("\n  Abstract elements:")
    elements = {
        "Background":  cfg["topic_keywords"],
        "Objective":   ["aim", "objective", "to systematically", "to examine"],
        "Methods":     ["search", "database", "review", "screen"],
        "Results":     ["found", "showed", "across", "result", "studies"],
        "Conclusions": ["suggest", "conclude", "evidence", "implication", "indicate"],
        "Keywords":    ["Keywords", "keyword"],
    }
    for elem, kws in elements.items():
        present = any(kw.lower() in abstract_text.lower() for kw in kws)
        print(f"    {'[OK]' if present else '[MISSING]'} {elem}")


# ── Part 4 ───────────────────────────────────────────────────────────

def check_part4_citations(doc, all_text, cfg):
    print("\n[PART 4] Citations & References")
    print("-" * 65)

    # et al. format
    et_al_bad = re.findall(r'et al[^.\s,;)\-]', all_text)
    print(f"  'et al' without period: {len(et_al_bad)} "
          f"{'[OK]' if not et_al_bad else '[WARN]'}")
    for s in et_al_bad[:3]:
        print(f"    Sample: ...{s}...")

    # In-text citations (both formats)
    intext_keys = extract_author_year_keys(all_text)
    print(f"\n  In-text citation keys (unique): {len(intext_keys)}")

    # Reference list
    ref_entries = extract_references(doc)
    ref_keys    = extract_ref_list_keys(ref_entries)
    print(f"  Reference list entries: {len(ref_entries)}")
    print(f"  Reference list keys (parsed): {len(ref_keys)}")

    doi_count = len(re.findall(r'https://doi\.org/', all_text))
    old_doi   = len(re.findall(r'(?<!https://)doi\.org/', all_text))
    print(f"  DOI links (https://doi.org/): {doi_count}")
    if old_doi:
        print(f"  [WARN] Old-style DOI (doi.org/ without https): {old_doi}")

    # Orphan citation check
    print("\n  Orphan citation check:")
    if intext_keys and ref_keys:
        in_not_ref = intext_keys - ref_keys
        ref_not_in = ref_keys - intext_keys

        # Filter noise: remove keys where author matches partially (regex limitations)
        # Only report if same (author, year) exact match missing
        if in_not_ref:
            print(f"  [WARN] Cited in text but NOT in reference list ({len(in_not_ref)}):")
            for key in sorted(in_not_ref)[:10]:
                print(f"    {key[0]} ({key[1]})")
            if len(in_not_ref) > 10:
                print(f"    ... and {len(in_not_ref)-10} more")
        else:
            print("  [OK] All in-text citations appear in reference list")

        if ref_not_in:
            print(f"  [NOTE] In reference list but not detected in text ({len(ref_not_in)}) "
                  f"— may be regex under-capture, verify manually:")
            for key in sorted(ref_not_in)[:5]:
                print(f"    {key[0]} ({key[1]})")
        else:
            print("  [OK] All reference list entries appear to be cited in text")
    else:
        print("  [SKIP] Insufficient data for orphan check")


# ── Part 5 ───────────────────────────────────────────────────────────

def check_part5_symbols(doc, all_text, cfg):
    print("\n[PART 5] Statistical Symbols & Formatting")
    print("-" * 65)

    # Remove DOI paths to avoid false positives in all regex below
    clean_text = re.sub(r'https?://\S+', ' __URL__ ', all_text)

    stat_patterns = {
        'p values (p =/</>)': r'\bp\s*[=<>]',
        'F statistics':        r'\bF\s*[=(]',
        't statistics':        r'\bt\s*[=(]',
        'effect sizes (d =)':  r'\bd\s*=',
        'k (studies)':         r'\bk\s*=',
        'n (subgroup)':        r'\bn\s*=',
        'N (total)':           r'\bN\s*=',
        'r (correlation)':     r'\br\s*=',
    }
    print("  Statistical symbol counts (verify italic in Word):")
    for sym, pat in stat_patterns.items():
        count = len(re.findall(pat, clean_text))
        if count > 0:
            snippets = context_snippet(clean_text, pat, n=2, window=30)
            print(f"    {sym}: {count}")
            for s in snippets:
                print(f"      e.g. {s}")
        else:
            print(f"    {sym}: 0")

    # APA number format checks
    print("\n  APA number format checks:")
    # p-values should NOT have leading zero in APA 7 (p = .034, not p = 0.034)
    # But d, r, R² etc. should have leading zero (d = 0.52, not d = .52)
    pval_with_zero = re.findall(r'\bp\s*[=<>]\s*0\.\d+', clean_text)
    if pval_with_zero:
        print(f"  [NOTE] p-values with leading zero ({len(pval_with_zero)}) — "
              f"APA 7 omits leading zero for p: use 'p = .034' not 'p = 0.034'")
        for s in pval_with_zero[:2]:
            print(f"    {s}")
    else:
        print("  [OK] p-value leading-zero format (none found, or already APA-correct)")

    # Effect sizes > 3 are suspicious (may be unstandardised or typo)
    large_d = re.findall(r'\bd\s*=\s*([3-9]\.\d+|\d{2,}\.\d+)', clean_text)
    if large_d:
        print(f"  [WARN] Unusually large effect size (d > 3): {large_d[:5]}")

    # En-dash in year/page ranges — exclude DOIs (already cleaned)
    # Match digit-hyphen-digit ranges that are NOT inside words (N-back, pre-post etc.)
    # Specifically: 4-digit year ranges and 2-3 digit page ranges
    year_hyphen = re.findall(r'\b((?:19|20)\d{2})-(?:19|20)\d{2}\b', clean_text)
    page_hyphen = re.findall(r'\b(\d{2,4})-(\d{2,4})\b(?!\s*\))', clean_text)
    # Filter out known compound words
    page_hyphen = [(a, b) for a, b in page_hyphen
                   if not re.match(r'^[A-Za-z]', a) and len(a) >= 2 and len(b) >= 2]
    if year_hyphen:
        print(f"\n  [WARN] Hyphen in year ranges (use en-dash –): {year_hyphen[:5]}")
    if page_hyphen:
        print(f"  [WARN] Possible hyphen in page ranges (use en-dash –): "
              f"{page_hyphen[:5]}")
    if not year_hyphen and not page_hyphen:
        print("\n  [OK] No hyphen-in-range issues detected")


# ── Part 6 ───────────────────────────────────────────────────────────

def check_part6_content(doc, all_text, cfg):
    print("\n[PART 6] Critical Content Checks")
    print("-" * 65)

    for item, pattern in cfg["required_statements"].items():
        present = bool(re.search(pattern, all_text, re.IGNORECASE))
        print(f"  {'[OK]' if present else '[MISSING]'} {item}")


# ── Part 7 ───────────────────────────────────────────────────────────

def check_part7_consistency(doc, all_text, cfg):
    print("\n[PART 7] Internal Consistency")
    print("-" * 65)

    n_inc = cfg["n_included"]

    # Study count
    count = len(re.findall(rf'\b{n_inc}\b', all_text))
    print(f"  Study count ({n_inc}) appearances: {count} "
          f"{'[OK]' if count >= 5 else '[CHECK: few mentions]'}")

    # Word count (body only: exclude abstract and references)
    body_text = extract_section(
        doc,
        start_marker="1. Introduction",
        stop_markers=["References", "Declarations"]
    )
    # fallback: try without numbering
    if len(body_text.split()) < 100:
        body_text = extract_section(doc, "Introduction", ["References", "Declarations"])
    body_wc = len(body_text.split())
    limit   = cfg.get("body_word_limit")
    limit_str = f"/ limit {limit:,}" if limit else ""
    limit_flag = (f" [WARN: exceeds {limit:,}]" if limit and body_wc > limit else
                  " [OK]" if limit else "")
    print(f"\n  Body word count (excl. Abstract & References): {body_wc:,}{limit_str}{limit_flag}")

    # Over-strong language
    print("\n  Language strength check:")
    for word in cfg["overstrong_language"]:
        count = len(re.findall(rf'\b{re.escape(word)}\b', all_text, re.IGNORECASE))
        neg   = bool(re.search(rf'\bin{re.escape(word)}\b', all_text, re.IGNORECASE))
        flag  = " [CHECK: consider softening]" if count > 0 and not neg else ""
        print(f"    '{word}': {count}{flag}")

    # Figure cross-check: references in text vs captions
    print("\n  Figure cross-check:")
    fig_refs = set(re.findall(r'\bFigure\s+(\d+)\b', all_text))
    # Search captions anywhere in text (not line-anchored)
    fig_caps = set(re.findall(r'Figure\s+(\d+)\.', all_text))
    print(f"    Referenced in text: {sorted(fig_refs)}")
    print(f"    Captions found:     {sorted(fig_caps)}")
    missing_caps = fig_refs - fig_caps
    missing_refs = fig_caps - fig_refs
    if missing_caps:
        print(f"    [WARN] Referenced but no caption: Figure {sorted(missing_caps)}")
    if missing_refs:
        print(f"    [NOTE] Caption exists but not referenced: Figure {sorted(missing_refs)}")
    if not missing_caps and not missing_refs:
        print(f"    [OK] Figure references and captions match")

    # Table cross-check
    print("\n  Table cross-check:")
    tbl_refs = set(re.findall(r'\bTable\s+(\d+)\b', all_text))
    tbl_caps = set(re.findall(r'Table\s+(\d+)\.', all_text))
    print(f"    Referenced in text: {sorted(tbl_refs)}")
    print(f"    Captions found:     {sorted(tbl_caps)}")
    print(f"    Actual tables in doc: {len(doc.tables)}")
    missing_tcaps = tbl_refs - tbl_caps
    missing_trefs = tbl_caps - tbl_refs
    if missing_tcaps:
        print(f"    [WARN] Referenced but no caption: Table {sorted(missing_tcaps)}")
    if missing_trefs:
        print(f"    [NOTE] Caption exists but not referenced: Table {sorted(missing_trefs)}")
    if not missing_tcaps and not missing_trefs:
        print(f"    [OK] Table references and captions match")


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    cfg = CONFIG
    fname = cfg["doc_path"].replace("\\", "/").split("/")[-1]
    print("=" * 65)
    print(f"PRE-SUBMISSION CHECKLIST v2: {cfg['paper_id']}")
    print(f"Topic:    {cfg['paper_topic']}")
    print(f"Document: {fname}")
    print("=" * 65)

    doc, all_text = load_doc(cfg["doc_path"])

    check_part1_completeness(doc, all_text, cfg)
    check_part2_numbers(doc, all_text, cfg)
    check_part3_abstract(doc, all_text, cfg)
    check_part4_citations(doc, all_text, cfg)
    check_part5_symbols(doc, all_text, cfg)
    check_part6_content(doc, all_text, cfg)
    check_part7_consistency(doc, all_text, cfg)

    print("\n" + "=" * 65)
    print("END OF REPORT")
    print("=" * 65)
    print("\n[NEXT STEPS]")
    print("  1. Fix all [WARN] and [MISSING] items")
    print("  2. Manually verify italic formatting for p, k, n, N, d, F, t in Word")
    print("  3. Spot-check orphan citations — regex may under-capture some keys")
    print("  4. Verify journal name italics and DOI links in reference list")
    print("  5. Re-run after fixes to confirm all clear")


if __name__ == "__main__":
    main()
