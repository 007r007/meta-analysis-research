"""
通用投稿前完整性审阅脚本
Universal Pre-Submission Checklist for Systematic Review Papers
==============================================================
适用范围：叙述性系统综述 / 元分析 论文(.docx格式)
基于 Paper2 审阅框架(presubmit_check_part1-6.py)重构，Paper2-specific内容已参数化

使用方法：
  1. 修改下方 CONFIG 区域中的论文专属参数
  2. 运行: python presubmit_check_universal.py

作者：CC (Claude Code)
创建：2026-04-06
适用项目：paper-01, paper-02, paper-03...
"""

import sys
import re
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

# ═══════════════════════════════════════════════════════════════════
# ★ CONFIG — 每篇论文修改这里 ★
# ═══════════════════════════════════════════════════════════════════
CONFIG = {
    # 文档路径（绝对路径）
    "doc_path": r"E:\Meta-analysis writing project\projects\paper-01\06-write\paper01_draft_v1.docx",

    # 论文标识（用于报告标题）
    "paper_id": "Paper-01",
    "paper_topic": "WM Training in Healthy Older Adults",

    # PRISMA 关键数字（从筛选Excel核实后填入）
    "n_total_records":  4168,   # 去重后总记录数
    "n_fulltext":        445,   # 进入全文筛选数
    "n_included":         56,   # 最终纳入数
    "n_excluded_stage12": 3723, # 标题摘要阶段排除数

    # 使用的数据库（list）
    "databases": ["PubMed", "PsycINFO", "Web of Science", "Scopus"],

    # 预期表格数量
    "expected_tables": 3,

    # 预期章节（按论文实际结构调整）
    "expected_sections": [
        "Abstract", "Introduction", "Methods",
        "Results", "Discussion", "Conclusion", "References"
    ],

    # 论文主题相关词（用于摘要背景检查）
    "topic_keywords": ["working memory", "older adult", "cognitive training"],

    # 内部代码格式（筛选阶段用的临时标记，投稿前应清除）
    # Paper1用seq编号，Paper2用[FTxx]；在草稿正文中如有残留会报警
    # 填None则跳过此检查
    "internal_code_pattern": None,   # e.g. r'\[FT\d{2}\]' or r'\bseq\d+\b'

    # 摘要字数上限（根据目标期刊调整）
    "abstract_word_limit": 300,

    # 需要核查的关键声明（缺失时报警）
    "required_statements": {
        "PRISMA 2020":              r"PRISMA",
        "Registration/PROSPERO":    r"PROSPERO|pre.regist|not pre.regist",
        "SWiM/Narrative synthesis": r"SWiM|narrative synthesis",
        "RoB tool cited":           r"RoB 2|RoB2|ROBINS.I|Sterne|risk of bias tool",
        "Near/Far transfer defined":r"near transfer|far transfer",
        "Limitations section":      r"[Ll]imitation",
        "Publication bias":         r"publication bias",
        "Conflict of interest":     r"conflict of interest|competing interest",
        "Funding statement":        r"funding|grant|no funding|no financial",
        "Data availability":        r"data avail|supplementary|supporting information",
        "Ethics/IRB":               r"ethics|IRB|institutional review|ethical approval",
    },

    # 语言过强词（建议软化）
    "overstrong_language": [
        "conclusive",
        "clearly demonstrates",
        "proves",
        "definitively",
        "unambiguously",
    ],
}
# ═══════════════════════════════════════════════════════════════════

def load_doc(path):
    doc = Document(path)
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text
    return doc, all_text


def check_part1_completeness(doc, all_text, cfg):
    print("\n[PART 1] Content Completeness")
    print("-" * 65)

    # 1a. Internal codes residual
    pattern = cfg.get("internal_code_pattern")
    if pattern:
        codes = re.findall(pattern, all_text)
        if codes:
            print(f"  [WARN] Internal codes still present ({len(codes)}): {codes[:5]}")
        else:
            print(f"  [OK] No internal codes found (pattern: {pattern})")
    else:
        print("  [SKIP] internal_code_pattern not set")

    # 1b. Generic placeholders
    print("\n  Placeholder check:")
    placeholders = ['XXX', 'TODO', 'TBD', '[INSERT', 'PLACEHOLDER', '???',
                    '[Author names to be added]', 'forthcoming', 'in press']
    found_any = False
    for ph in placeholders:
        if ph.lower() in all_text.lower():
            count = all_text.lower().count(ph.lower())
            print(f"    [WARN] '{ph}': {count} times")
            found_any = True
    if not found_any:
        print("    [OK] No placeholders found")

    # 1c. Section structure
    print("\n  Section structure:")
    for s in cfg["expected_sections"]:
        found = any(s in p.text for p in doc.paragraphs)
        print(f"    {'[OK]' if found else '[MISSING]'} {s}")

    # 1d. Tables
    n_tables = len(doc.tables)
    exp = cfg["expected_tables"]
    status = "[OK]" if n_tables == exp else f"[WARN: expected {exp}]"
    print(f"\n  Tables: {n_tables} found {status}")
    for i, tbl in enumerate(doc.tables):
        print(f"    Table {i+1}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")


def check_part2_numbers(doc, all_text, cfg):
    print("\n[PART 2] Key Numbers Verification")
    print("-" * 65)

    n_total   = cfg["n_total_records"]
    n_ft      = cfg["n_fulltext"]
    n_inc     = cfg["n_included"]
    n_excl12  = cfg["n_excluded_stage12"]

    # Verify PRISMA arithmetic
    calc_ft   = n_total - n_excl12
    calc_excl_ft = n_ft - n_inc

    print(f"  PRISMA arithmetic:")
    print(f"    Total({n_total}) - ExclStage12({n_excl12}) = {calc_ft} "
          f"{'== fulltext OK' if calc_ft == n_ft else f'!= fulltext({n_ft}) [MISMATCH]'}")
    print(f"    Fulltext({n_ft}) - Included({n_inc}) = {calc_excl_ft} "
          f"(= full-text excluded)")

    print(f"\n  Key numbers in text:")
    for label, num in [
        (f"Total records ({n_total})", str(n_total).replace(',', '')),
        (f"Fulltext ({n_ft})", str(n_ft)),
        (f"Included ({n_inc})", str(n_inc)),
    ]:
        pattern = num.replace(',', '[,.]?')
        count = len(re.findall(rf'\b{pattern}\b', all_text))
        print(f"    {label}: {count} mentions {'[OK]' if count >= 2 else '[CHECK: few mentions]'}")

    print(f"\n  Databases:")
    for db in cfg["databases"]:
        n = len(re.findall(re.escape(db), all_text))
        print(f"    {db}: {n} times {'[OK]' if n > 0 else '[MISSING]'}")


def check_part3_abstract(doc, all_text, cfg):
    print("\n[PART 3] Abstract")
    print("-" * 65)

    abstract_text = ""
    in_abstract = False
    for para in doc.paragraphs:
        if para.text.strip() == "Abstract":
            in_abstract = True
            continue
        if in_abstract:
            if para.text.strip().startswith("1.") or para.text.strip() in ("1. Introduction", "Keywords"):
                # Don't break on Keywords - keep collecting
                pass
            if para.text.strip().startswith("1."):
                break
            abstract_text += para.text + " "

    wc = len(abstract_text.split())
    limit = cfg["abstract_word_limit"]
    print(f"  Word count: {wc} {'[OK]' if wc <= limit else f'[WARN: exceeds {limit}]'}")

    print(f"\n  Abstract elements:")
    elements = {
        "Background":   cfg["topic_keywords"],
        "Objective":    ["aim", "objective", "to systematically", "to examine"],
        "Methods":      ["search", "database", "review", "screen"],
        "Results":      ["found", "showed", "across", "result", "studies"],
        "Conclusions":  ["suggest", "conclude", "evidence", "implication", "indicate"],
        "Keywords":     ["Keywords", "keyword"],
    }
    for elem, keywords in elements.items():
        present = any(kw.lower() in abstract_text.lower() for kw in keywords)
        print(f"    {'[OK]' if present else '[MISSING]'} {elem}")


def check_part4_citations(doc, all_text, cfg):
    print("\n[PART 4] Citations & References")
    print("-" * 65)

    # et al. check
    et_al_bad = re.findall(r'et al[^.\s,)]', all_text)
    print(f"  'et al' without period: {len(et_al_bad)} "
          f"{'[OK]' if not et_al_bad else '[WARN]'}")
    for sample in et_al_bad[:3]:
        print(f"    Sample: ...{sample}...")

    # Citation count
    citation_pattern = (r'\(([A-Z][a-zA-Z\u00c0-\u024f\-]+(?:\s+(?:et al\.|&|and)'
                        r'\s+[A-Z][a-z]+)?)(?:,)?\s+(\d{4}[a-z]?)\)')
    citations = re.findall(citation_pattern, all_text)
    print(f"  In-text citations: {len(citations)} total, {len(set(citations))} unique")

    # Reference entries
    ref_start = False
    ref_entries = []
    for para in doc.paragraphs:
        if para.text.strip() == "References":
            ref_start = True
            continue
        if ref_start and para.text.strip() and para.text[0].isupper():
            ref_entries.append(para.text.strip())
    print(f"  Reference list entries: {len(ref_entries)}")

    doi_count = len(re.findall(r'https://doi\.org|doi\.org', all_text))
    print(f"  DOI links: {doi_count}")

    if len(ref_entries) > 0:
        ratio = len(citations) / len(ref_entries)
        if ratio < 0.3:
            print(f"  [CHECK] Citation/Reference ratio low ({ratio:.1f}) — "
                  f"possible mismatch or regex under-capture")


def check_part5_symbols(doc, all_text, cfg):
    print("\n[PART 5] Statistical Symbols & Formatting")
    print("-" * 65)

    stat_symbols = {
        'p values (p =/</>)':  r'\bp\s*[=<>]',
        'F statistics':         r'\bF\s*[=(]',
        't statistics':         r'\bt\s*[=(]',
        'effect sizes (d =)':   r'\bd\s*=',
        'k (studies)':          r'\bk\s*=',
        'n (subgroup)':         r'\bn\s*=',
        'N (total sample)':     r'\bN\s*=',
    }
    for sym, pat in stat_symbols.items():
        count = len(re.findall(pat, all_text))
        print(f"  {sym}: {count}")
    print("  [NOTE] Verify p, k, n, N, d, F, t are italic in Word")

    # En-dash vs hyphen in ranges
    hyphen_ranges = re.findall(r'\d+-\d{4}', all_text)  # e.g. 2003-2026 should be 2003–2026
    if hyphen_ranges:
        print(f"\n  [WARN] Possible hyphen in year ranges (should be en-dash –): "
              f"{hyphen_ranges[:5]}")
    else:
        print(f"\n  [OK] No hyphen-in-range issues detected")


def check_part6_content(doc, all_text, cfg):
    print("\n[PART 6] Critical Content Checks")
    print("-" * 65)

    for item, pattern in cfg["required_statements"].items():
        present = bool(re.search(pattern, all_text, re.IGNORECASE))
        print(f"  {'[OK]' if present else '[MISSING]'} {item}")


def check_part7_consistency(doc, all_text, cfg):
    print("\n[PART 7] Internal Consistency")
    print("-" * 65)

    n_inc = cfg["n_included"]

    # Study count consistency
    count = len(re.findall(rf'\b{n_inc}\b', all_text))
    print(f"  Study count ({n_inc}) appearances: {count} "
          f"{'[OK]' if count >= 5 else '[CHECK: few mentions]'}")

    # Over-strong language
    print(f"\n  Language strength check:")
    for word in cfg["overstrong_language"]:
        count = len(re.findall(rf'\b{re.escape(word)}\b', all_text, re.IGNORECASE))
        # "inconclusive" is fine; warn only on affirmative strong claims
        is_negative = bool(re.search(rf'\bin{re.escape(word)}\b', all_text, re.IGNORECASE))
        flag = ""
        if count > 0 and not is_negative:
            flag = " [CHECK: consider softening]"
        print(f"    '{word}': {count}{flag}")

    # Figure references vs figure captions
    fig_refs   = set(re.findall(r'\bFigure\s+(\d+)\b', all_text))
    fig_caps   = set(re.findall(r'^Figure\s+(\d+)\.', all_text, re.MULTILINE))
    print(f"\n  Figure references in text: {sorted(fig_refs)}")
    print(f"  Figure captions found: {sorted(fig_caps)}")
    if fig_refs != fig_caps:
        print(f"  [CHECK] Mismatch — referenced: {fig_refs}, captioned: {fig_caps}")
    else:
        print(f"  [OK] Figure references and captions match")

    # Table references
    tbl_refs  = set(re.findall(r'\bTable\s+(\d+)\b', all_text))
    print(f"\n  Table references in text: {sorted(tbl_refs)}")
    print(f"  Actual tables in doc: {len(doc.tables)}")


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════
def main():
    cfg = CONFIG
    print("=" * 65)
    print(f"PRE-SUBMISSION CHECKLIST: {cfg['paper_id']}")
    print(f"Topic: {cfg['paper_topic']}")
    print(f"Document: {cfg['doc_path'].split(chr(92))[-1]}")
    print("=" * 65)

    doc, all_text = load_doc(cfg["doc_path"])

    check_part1_completeness(doc, all_text, cfg)
    check_part2_numbers(doc, all_text, cfg)
    check_part3_abstract(doc, all_text, cfg)
    check_part4_citations(doc, all_text, cfg)
    check_part5_symbols(doc, all_text, cfg)
    check_part6_content(doc, all_text, cfg)
    check_part7_consistency(doc, all_text, cfg)

    print("\n" + "=" * 65)
    print("END OF REPORT")
    print("=" * 65)
    print("\n[NEXT STEPS]")
    print("  1. Fix all [WARN] and [MISSING] items above")
    print("  2. Manually verify italic formatting for statistical symbols in Word")
    print("  3. Check reference list visually for journal name italics and DOI links")
    print("  4. Re-run this script after fixes to confirm all clear")


if __name__ == "__main__":
    main()
